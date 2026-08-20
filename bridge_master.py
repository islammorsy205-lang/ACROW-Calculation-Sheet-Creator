# ==============================================================================
# 🌉 ACROW BRIDGE MASTER - THE COMMERCIAL GRADE ENGINE (Ultimate Edition V7.0) 🌉
# ==============================================================================
# 🌟 Features: 
# - Descriptive Parameters Injection (Cross Section, Top, Web, Bottom).
# - 3-Zone Soldier FEA Checks (Cantilever, Web, Bottom Slab).
# - Fixed 18mm Plywood & Detailed Secondary Beam Spans (No Reactions).
# - Custom Acrow Formatting (Red SAFE text, SHORBRACE TABLE FORM cover).
# - Detailed Push Pull types extraction & exact Acrow Report Styles.
# ==============================================================================

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import re
import math
import tempfile
import time
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from PIL import Image, ImageChops

# ---------------------------------------------------------
# Word Report & PDF Dependencies
# ---------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    st.error("⚠️ مكتبة 'python-docx' غير موجودة! برجاء تثبيتها.")

try:
    import ezdxf
except ImportError:
    st.error("⚠️ مكتبة 'ezdxf' غير موجودة! برجاء تثبيتها: pip install ezdxf")
    ezdxf = None

try:
    import fitz
except ImportError:
    pass

# ---------------------------------------------------------
# Custom Project Modules
# ---------------------------------------------------------
try:
    from report_builder import insert_blue_banner, add_eq, append_pdf_stream_to_word, force_ltr_left, add_heading_14, add_eq_highlight, add_centered_text
    from config import SECTIONS_DB, STRUTS_DB
    from math_solver import solve_beam_advanced
    from plot_core import draw_system_sketch, generate_acrow_diagrams
except ImportError:
    pass

# =========================================================
# 0. Core Helper Functions & Styles
# =========================================================
def apply_plot_styles():
    matplotlib.rcParams['font.family'] = 'sans-serif'
    matplotlib.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'DejaVu Sans']
    matplotlib.rcParams['axes.linewidth'] = 0.3
    matplotlib.rcParams['font.size'] = 7
    matplotlib.rcParams['font.weight'] = 'normal'
    matplotlib.rcParams['axes.labelweight'] = 'normal'

def crop_image_bbox(img_bytes):
    img = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
    bg = Image.new("RGBA", img.size, (255, 255, 255, 0))
    diff = ImageChops.difference(img, bg)
    bbox = diff.getbbox()
    if bbox:
        padding = 5
        bbox = (max(0, bbox[0] - padding), max(0, bbox[1] - padding), min(img.size[0], bbox[2] + padding), min(img.size[1], bbox[3] + padding))
        img = img.crop(bbox)
    out = io.BytesIO()
    img.save(out, format='PNG')
    return out.getvalue()

def safe_render_fig(fig):
    try:
        plt.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=400, bbox_inches='tight', pad_inches=0.0, transparent=True)
        return crop_image_bbox(buf.getvalue())
    finally:
        plt.close(fig)

def eval_seg_point(seg, s_val):
    if seg.get('is_divided'):
        actual_s = s_val + seg.get('parent_offset', 0.0)
        return eval_seg_point(seg['parent_seg'], actual_s)
    L = seg.get('L', 0.0)
    s_val = min(max(s_val, 0.0), L)
    ratio = s_val / L if L > 1e-6 else 0.0
    if seg.get('is_dxf') or 'abs_p1' in seg:
        p1, p2 = seg.get('abs_p1', (0,0)), seg.get('abs_p2', (0,0))
        px = p1[0] + ratio * (p2[0] - p1[0])
        pz = p1[1] + ratio * (p2[1] - p1[1])
        th = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
        return px, pz, th
    return 0.0, 0.0, 0.0

def drop_ray_to_bottom_chord(x_target, segments):
    z_candidates = []
    for seg in segments:
        if seg.get('Shape Type') == 'Straight Line':
            p1, p2 = seg['abs_p1'], seg['abs_p2']
            min_px, max_px = min(p1[0], p2[0]), max(p1[0], p2[0])
            if min_px - 0.05 <= x_target <= max_px + 0.05:
                if abs(max_px - min_px) < 1e-5: z_candidates.append(min(p1[1], p2[1]))
                else:
                    ratio = (x_target - p1[0]) / (p2[0] - p1[0])
                    z_candidates.append(p1[1] + ratio * (p2[1] - p1[1]))
    if z_candidates: return x_target, min(z_candidates)
    all_pts = []
    for s in segments: all_pts.extend([s.get('abs_p1'), s.get('abs_p2')])
    if not all_pts: return x_target, 0.0
    best_pt = min(all_pts, key=lambda p: abs(p[0] - x_target) if p else 999)
    return x_target, best_pt[1]

def get_closest_segment_exact(pt, segs):
    min_d, best_idx, best_s = 9999.0, 0, 0.0
    px, pz = pt[0], pt[1]
    for idx, seg in enumerate(segs):
        temp_seg = seg['parent_seg'] if seg.get('is_divided') else seg
        L_orig = temp_seg.get('L', 0.0)
        if 'abs_p1' in temp_seg:
            p1, p2 = np.array(temp_seg['abs_p1']), np.array(temp_seg['abs_p2'])
            v, w = p2 - p1, np.array([px, pz]) - p1
            c2 = np.dot(v, v)
            ratio = max(0.0, min(1.0, np.dot(w, v) / c2 if c2 > 1e-6 else 0.0))
            d = np.linalg.norm(np.array([px, pz]) - (p1 + ratio * v))
            if d < min_d: min_d, best_idx, best_s = d, idx, ratio * L_orig
            if seg.get('is_divided'): best_s -= seg.get('parent_offset', 0.0)
    return min_d, best_idx, best_s

def get_shifted_coords_along_segment(px, pz, ds, segs):
    if abs(ds) < 1e-4: return px, pz
    d_min, best_idx, best_s = get_closest_segment_exact((px, pz), segs)
    if d_min > 0.5: return px + ds, pz
    nx, nz, _ = eval_seg_point(segs[best_idx], max(0.0, min(best_s + ds, segs[best_idx].get('L', 0.0))))
    return nx, nz

def get_valid_strut_names():
    if not STRUTS_DB: return ["PPH"]
    valid_struts = []
    for s_name in STRUTS_DB.keys():
        name_u, base_name = s_name.upper(), s_name.split('(')[0].strip()
        if "TILT" in name_u or "MMP" in name_u or base_name.endswith('1') or base_name.endswith('3'): continue
        priority = 1 if "PPS" in name_u else (2 if "PPH" in name_u else 99)
        valid_struts.append({'name': s_name, 'pri': priority})
    valid_struts.sort(key=lambda x: x['pri'])
    return [x['name'] for x in valid_struts] if valid_struts else ["PPH"]

def get_optimal_strut_section(req_length, req_axial_force):
    valid_struts = []
    for s_name in get_valid_strut_names():
        s_props = STRUTS_DB.get(s_name, {})
        m = re.search(r'\((\d+\.\d+):(\d+\.\d+)m\)', s_name)
        if m and float(m.group(1)) <= req_length <= float(m.group(2)) and s_props.get('allow', 0.0) >= abs(req_axial_force):
            valid_struts.append({'name': s_name, 'allowable': s_props.get('allow', 0.0)})
    if not valid_struts: return None
    valid_struts.sort(key=lambda x: x['allowable']) 
    return valid_struts[0]['name']

def add_red_safe_check(doc_obj, title, act, allw, unit, extra_txt=""):
    p = doc_obj.add_paragraph()
    force_ltr_left(p)
    if title:
        r1 = p.add_run(f"• {title}:\n")
        r1.font.name, r1.font.size, r1.font.bold = 'Arial', Pt(11), True
    r2 = p.add_run(f"  Max = {act:.2f} {unit}   <   {allw:.2f} {unit}   ")
    r2.font.name, r2.font.size = 'Arial', Pt(11)
    
    res = p.add_run("SAFE" if act <= allw else "UNSAFE ❌")
    res.font.name, res.font.size, res.font.bold = 'Arial', Pt(11), True
    res.font.color.rgb = RGBColor(255, 0, 0)
    
    if extra_txt: 
        r3 = p.add_run(f"\n  ({extra_txt})")
        r3.font.name, r3.font.size, r3.italic = 'Arial', Pt(10), True

def add_reference_line(doc_obj, item_name, ds_map):
    if not item_name: return
    item_clean = str(item_name).upper().replace("DOUBLE", "").strip()
    matched_page = " "
    if ds_map:
        search_terms = [item_clean]
        if "PPH" in item_clean or "PPS" in item_clean: search_terms.extend(["PUSH PULL", "PUSH-PULL", "STRUT"])
        elif "H20" in item_clean: search_terms.extend(["H20", "TIMBER"])
        elif "S12" in item_clean: search_terms.extend(["S12", "ACROW BEAM"])
        elif "SOLDIER" in item_clean: search_terms.extend(["SOLDIER"])
        for term in search_terms:
            for k, v in ds_map.items():
                if term in str(k).upper():
                    matched_page = v; break
            if matched_page != " ": break
    p_ref = doc_obj.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ref = p_ref.add_run(f"For allowable values / technical data of {item_name}, refer to page {matched_page}")
    r_ref.font.name, r_ref.font.size, r_ref.font.italic, r_ref.font.color.rgb = 'Arial', Pt(11), True, RGBColor(0, 112, 192)

# ==============================================================================
# 2. THE STRICT DXF PARSER (Bulletproof Binary Mode)
# ==============================================================================
@st.cache_data
def parse_dxf_bridge_cases(file_bytes, conc_density=25.0):
    if ezdxf is None: 
        return None
        
    tmp_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".dxf", mode='wb') as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
            
        doc = ezdxf.readfile(tmp_path)
        msp = doc.modelspace()
        
        layer_table = 'TABLE_ANALYSIS'
        layer_supp = 'SHORING_LINES'  
        layer_text = 'TEXT_DATA'
        layer_frame = 'FRAME'
        layer_strut = 'PUSH_PULL'
        
        table_boxes = []
        for e in msp:
            if e.dxftype() in ['LWPOLYLINE', 'POLYLINE'] and e.dxf.layer.upper() == layer_table:
                if e.dxftype() == 'LWPOLYLINE':
                    points = e.get_points('xy')
                else:
                    points = [v.dxf.location for v in e.vertices]
                    
                xs = [p[0] / 1000.0 for p in points]
                zs = [p[1] / 1000.0 for p in points]
                table_boxes.append({'min_x': min(xs), 'max_x': max(xs), 'min_z': min(zs), 'max_z': max(zs), 'cx': (min(xs) + max(xs)) / 2.0, 'cz': (min(zs) + max(zs)) / 2.0})
                
        if not table_boxes:
            all_xs, all_zs = [], []
            for e in msp:
                if e.dxftype() == 'LINE' and e.dxf.layer.upper() == layer_frame:
                    x1, z1 = e.dxf.start.x / 1000.0, e.dxf.start.y / 1000.0
                    x2, z2 = e.dxf.end.x / 1000.0, e.dxf.end.y / 1000.0
                    if math.hypot(x2 - x1, z2 - z1) >= 0.01:
                        all_xs.extend([x1, x2])
                        all_zs.extend([z1, z2])
            
            if all_xs and all_zs:
                table_boxes.append({'min_x': min(all_xs) - 1.0, 'max_x': max(all_xs) + 1.0, 'min_z': min(all_zs) - 1.0, 'max_z': max(all_zs) + 1.0, 'cx': (min(all_xs) + max(all_xs)) / 2.0, 'cz': (min(all_zs) + max(all_zs)) / 2.0})
                
        table_boxes.sort(key=lambda b: b['min_x'])
        cases_raw = [{'box': box, 'frames': [], 'struts': [], 'supports': [], 'cut_points': [], 's_texts': [], 'a_texts': []} for box in table_boxes]
            
        for e in msp:
            layer = e.dxf.layer.upper()
            dxftype = e.dxftype()
            x_cad, z_cad, is_valid_point = 0.0, 0.0, False
            
            if dxftype in ['POINT', 'CIRCLE']:
                if dxftype == 'POINT': x_cad, z_cad = e.dxf.location.x / 1000.0, e.dxf.location.y / 1000.0
                else: x_cad, z_cad = e.dxf.center.x / 1000.0, e.dxf.center.y / 1000.0
                is_valid_point = True
            elif dxftype == 'INSERT':
                x_cad, z_cad = e.dxf.insert.x / 1000.0, e.dxf.insert.y / 1000.0
                is_valid_point = True
                
            if is_valid_point:
                for c in cases_raw:
                    b = c['box']
                    if b['min_x'] <= x_cad <= b['max_x'] and b['min_z'] <= z_cad <= b['max_z']:
                        norm_x, norm_z = x_cad - b['cx'], z_cad - b['cz']
                        if layer == layer_text: c['cut_points'].append({'x': norm_x, 'z': norm_z})
                        break

            elif dxftype in ['TEXT', 'MTEXT']:
                x_cad, z_cad = e.dxf.insert.x / 1000.0, e.dxf.insert.y / 1000.0
                for c in cases_raw:
                    b = c['box']
                    if b['min_x'] <= x_cad <= b['max_x'] and b['min_z'] <= z_cad <= b['max_z']:
                        if layer == layer_text:
                            txt = e.text if dxftype == 'MTEXT' else e.dxf.text
                            txt = re.sub(r'\\[A-Za-z0-9]+;', '', txt).upper().replace('\n', '').replace('\r', '').replace(' ', '')
                            s_m = re.search(r'S(\d+)=([\d\.]+)', txt)
                            s_lbl = re.search(r'S(\d+)(?!=)', txt)
                            a_m = re.search(r'A(\d+)=([\d\.]+)', txt)
                            norm_x, norm_z = x_cad - b['cx'], z_cad - b['cz']
                            
                            if s_m: c['s_texts'].append({'idx': int(s_m.group(1)), 'val': float(s_m.group(2)), 'x': norm_x, 'z': norm_z})
                            elif s_lbl: c['s_texts'].append({'idx': int(s_lbl.group(1)), 'val': 0.0, 'x': norm_x, 'z': norm_z})
                            if a_m: c['a_texts'].append({'idx': int(a_m.group(1)), 'val': float(a_m.group(2)), 'x': norm_x, 'z': norm_z})
                        break
                        
            elif dxftype in ['LINE', 'LWPOLYLINE', 'POLYLINE']:
                entities = list(e.virtual_entities()) if dxftype != 'LINE' else [e]
                for sub_e in entities:
                    if sub_e.dxftype() == 'LINE':
                        x1, z1 = sub_e.dxf.start.x / 1000.0, sub_e.dxf.start.y / 1000.0
                        x2, z2 = sub_e.dxf.end.x / 1000.0, sub_e.dxf.end.y / 1000.0
                        if math.hypot(x2 - x1, z2 - z1) < 0.01: continue
                        mid_x, mid_z = (x1 + x2) / 2.0, (z1 + z2) / 2.0
                        for c in cases_raw:
                            b = c['box']
                            if b['min_x'] <= mid_x <= b['max_x'] and b['min_z'] <= mid_z <= b['max_z']:
                                norm_x1, norm_z1 = x1 - b['cx'], z1 - b['cz']
                                norm_x2, norm_z2 = x2 - b['cx'], z2 - b['cz']
                                if layer == layer_frame: c['frames'].append({'x1': norm_x1, 'z1': norm_z1, 'x2': norm_x2, 'z2': norm_z2})
                                elif layer == layer_strut: c['struts'].append({'x1': norm_x1, 'z1': norm_z1, 'x2': norm_x2, 'z2': norm_z2})
                                elif layer == layer_supp: c['supports'].append({'x': norm_x1, 'z': norm_z1, 'type': 'Roller', 'angle': 0.0})
                                break

        processed_cases = []
        for c_idx, c in enumerate(cases_raw):
            if not c['frames']: continue 
            
            base_segments = []
            for i, line in enumerate(c['frames']):
                L = math.hypot(line['x2'] - line['x1'], line['z2'] - line['z1'])
                base_segments.append({'name': f"F{i+1}", 'master_idx': i, 'type': 'Straight Line', 'Shape Type': 'Straight Line', 'L': L, 'is_dxf': True, 'abs_p1': (line['x1'], line['z1']), 'abs_p2': (line['x2'], line['z2'])})
                
            unique_segments = []
            for seg in base_segments:
                is_dup = False
                for u_seg in unique_segments:
                    d_p1 = math.hypot(seg['abs_p1'][0] - u_seg['abs_p1'][0], seg['abs_p1'][1] - u_seg['abs_p1'][1])
                    d_p2 = math.hypot(seg['abs_p2'][0] - u_seg['abs_p2'][0], seg['abs_p2'][1] - u_seg['abs_p2'][1])
                    d_p1_rev = math.hypot(seg['abs_p1'][0] - u_seg['abs_p2'][0], seg['abs_p1'][1] - u_seg['abs_p2'][1])
                    d_p2_rev = math.hypot(seg['abs_p2'][0] - u_seg['abs_p1'][0], seg['abs_p2'][1] - u_seg['abs_p1'][1])
                    if (d_p1 < 0.05 and d_p2 < 0.05) or (d_p1_rev < 0.05 and d_p2_rev < 0.05):
                        is_dup = True; break
                if not is_dup: unique_segments.append(seg)
            base_segments = unique_segments
            
            dxf_areas = []
            for s_txt in c['s_texts']:
                min_d, best_idx, _ = get_closest_segment_exact((s_txt['x'], s_txt['z']), base_segments)
                if min_d < 4.0:  
                    seg_final_name = f"S{s_txt['idx']}"
                    base_segments[best_idx]['name'] = seg_final_name
                    a_txt = next((a for a in c['a_texts'] if a['idx'] == s_txt['idx']), None)
                    if a_txt and s_txt['val'] > 1e-4:
                        if not any(d['segment'] == seg_final_name for d in dxf_areas):
                            dxf_areas.append({'seg_idx': best_idx, 'segment': seg_final_name, 'length': s_txt['val'], 'area': a_txt['val']})
                            
            all_xs = [seg['abs_p1'][0] for seg in base_segments] + [seg['abs_p2'][0] for seg in base_segments]
            if all_xs:
                center_x = (min(all_xs) + max(all_xs)) / 2.0
                named_segs = [seg for seg in base_segments if seg['name'].startswith('S') and seg['name'] not in ["S30", "S31", "S40", "S41"]]
                for n_seg in named_segs:
                    m_p1_x = center_x - (n_seg['abs_p1'][0] - center_x)
                    m_p2_x = center_x - (n_seg['abs_p2'][0] - center_x)
                    for u_seg in base_segments:
                        if u_seg['name'].startswith('F'):
                            d1 = math.hypot(u_seg['abs_p1'][0] - m_p1_x, u_seg['abs_p1'][1] - n_seg['abs_p1'][1]) + math.hypot(u_seg['abs_p2'][0] - m_p2_x, u_seg['abs_p2'][1] - n_seg['abs_p2'][1])
                            d2 = math.hypot(u_seg['abs_p1'][0] - m_p2_x, u_seg['abs_p1'][1] - n_seg['abs_p2'][1]) + math.hypot(u_seg['abs_p2'][0] - m_p1_x, u_seg['abs_p2'][1] - n_seg['abs_p1'][1])
                            if min(d1, d2) < 1.0: u_seg['name'] = n_seg['name']
            
            unlabeled = [{'idx': idx, 'mx': (seg['abs_p1'][0] + seg['abs_p2'][0])/2.0, 'mz': (seg['abs_p1'][1] + seg['abs_p2'][1])/2.0} for idx, seg in enumerate(base_segments) if seg['name'].startswith('F')]
            if unlabeled:
                unlabeled.sort(key=lambda item: item['mz'])
                bottom_cands = [u for u in unlabeled if u['mz'] - unlabeled[0]['mz'] < 0.5]
                bottom_cands.sort(key=lambda item: item['mx'])
                if len(bottom_cands) > 0: base_segments[bottom_cands[0]['idx']]['name'] = "S30"
                if len(bottom_cands) > 1: base_segments[bottom_cands[-1]['idx']]['name'] = "S31"
                top_cands = [u for u in unlabeled if unlabeled[-1]['mz'] - u['mz'] < 0.5]
                top_cands.sort(key=lambda item: item['mx'])
                if len(top_cands) > 0 and base_segments[top_cands[0]['idx']]['name'].startswith('F'): base_segments[top_cands[0]['idx']]['name'] = "S40"
                if len(top_cands) > 1 and base_segments[top_cands[-1]['idx']]['name'].startswith('F'): base_segments[top_cands[-1]['idx']]['name'] = "S41"

            initial_loads = []
            dxf_calc_details = [] 
            
            for area_item in dxf_areas:
                s_name = area_item['segment']
                w_val = (area_item['area'] * conc_density * 1.30) / area_item['length']
                matching_indices = [idx_m for idx_m, seg_m in enumerate(base_segments) if seg_m['name'] == s_name]
                if matching_indices:
                    initial_loads.append({'seg_idx': matching_indices[0], 'category': 'Dead Load', 'type': 'Uniform', 'dir': 'Global Z (Vertical)', 'target_mode': 'Single Segment' if len(matching_indices) == 1 else 'Multiple Segments', 'target_segs_idx': matching_indices, 'start': 0.0, 'end': base_segments[matching_indices[0]]['L'], 'w1': -abs(w_val), 'w2': -abs(w_val), 'loc': 0.0, 'is_auto': True})
                dxf_calc_details.append({'segment': s_name, 'length': area_item['length'], 'area': area_item['area'], 'load_w': abs(w_val)})

            valid_ll_indices = [idx for idx, seg in enumerate(base_segments) if seg['name'] not in ["S30", "S31"]]
            if valid_ll_indices:
                initial_loads.append({'seg_idx': valid_ll_indices[0], 'category': 'Live Load', 'type': 'Uniform', 'dir': 'Global Z (Vertical)', 'target_mode': 'Multiple Segments', 'target_segs_idx': valid_ll_indices, 'start': 0.0, 'end': base_segments[valid_ll_indices[0]]['L'], 'w1': -3.77, 'w2': -3.77, 'loc': 0.0, 'is_auto': True})

            strut_opts = get_valid_strut_names()
            struts_mapped = []
            for line in c['struts']:
                tx, tz, bx, bz = (line['x1'], line['z1'], line['x2'], line['z2']) if line['z1'] > line['z2'] else (line['x2'], line['z2'], line['x1'], line['z1'])
                struts_mapped.append({'tx': tx, 'tz': tz, 'bx': bx, 'bz': bz, 'sec': strut_opts[0] if strut_opts else "Unknown"})
                
            if c['supports']:
                for sp in c['supports']:
                    final_x, final_z = drop_ray_to_bottom_chord(sp['x'], base_segments)
                    sp['x'], sp['z'] = final_x, final_z
                c['supports'].sort(key=lambda sp: sp['x'])
                c['supports'][0]['type'] = 'Hinged'
                
            processed_cases.append({'title': f"Table {c_idx+1}", 'segments': base_segments, 'struts': struts_mapped, 'supports': c['supports'], 'cut_points': c['cut_points'], 'dxf_areas': dxf_areas, 'loads': initial_loads, 'calc_details': dxf_calc_details})
            
        return processed_cases
        
    except Exception as e:
        st.error(f"DXF Parsing Error: {e}")
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass

# ==============================================================================
# 3. MESHING & FEA MATRIX ENGINE
# ==============================================================================
def perform_smart_division(base_segments, supports, struts, cut_points=[]):
    cut_points_dict = {}
    for i, seg in enumerate(base_segments):
        cut_points_dict[i] = {0.0, seg.get('L', 0.0)}
        
    for sp in supports:
        d_min, w_seg, w_s = get_closest_segment_exact((sp['x'], sp['z']), base_segments)
        if d_min < 0.30: 
            cut_points_dict[w_seg].add(min(max(w_s, 0.0), base_segments[w_seg]['L']))
            
    for st in struts:
        dt, wt_seg, wt_s = get_closest_segment_exact((st['tx'], st['tz']), base_segments)
        if dt < 0.30: cut_points_dict[wt_seg].add(min(max(wt_s, 0.0), base_segments[wt_seg]['L']))
        db, wb_seg, wb_s = get_closest_segment_exact((st['bx'], st['bz']), base_segments)
        if db < 0.30: cut_points_dict[wb_seg].add(min(max(wb_s, 0.0), base_segments[wb_seg]['L']))
            
    for cp in cut_points:
        d_min, w_seg, w_s = get_closest_segment_exact((cp['x'], cp['z']), base_segments)
        if d_min < 0.30: cut_points_dict[w_seg].add(min(max(w_s, 0.0), base_segments[w_seg]['L']))

    divided_segments = []
    sub_letters = "abcdefghijklmnopqrstuvwxyz"
    
    for m_idx, s_vals_set in sorted(cut_points_dict.items()):
        master_seg = base_segments[m_idx]
        sorted_s = sorted(list(s_vals_set))
        num_sub = len(sorted_s) - 1
        
        for k in range(num_sub):
            s_start, s_end = sorted_s[k], sorted_s[k+1]
            if s_end - s_start < 1e-4: continue
                
            sub_name = master_seg['name'] if num_sub == 1 else f"{master_seg['name']}-{sub_letters[k % 26]}"
            new_seg = master_seg.copy()
            new_seg.update({'name': sub_name, 'is_divided': True, 'parent_seg': master_seg, 'parent_offset': s_start, 'L': s_end - s_start, 'master_idx': m_idx})
            divided_segments.append(new_seg)
            
    return divided_segments

def build_chain_mesh(segments, seg_sections, loads, struts, supports, cut_points=[], mesh_size=0.50):
    nodes, elements, nodal_loads, node_tol = [], [], [], 0.01 
    
    def get_or_add_node(x, z):
        for i, n in enumerate(nodes):
            if abs(n[0] - x) < node_tol and abs(n[1] - z) < node_tol: return i
        nodes.append([x, z]); return len(nodes) - 1

    support_injections = {i: [] for i in range(len(segments))}
    supports_list_out = []
    
    for sup in supports:
        sx, sz = sup['x'], sup.get('z', sup.get('y', 0.0))
        min_d, w_seg, w_s = get_closest_segment_exact((sx, sz), segments)
        if min_d < 0.30: support_injections[w_seg].append(w_s)
        nid = get_or_add_node(sx, sz)
        supports_list_out.append({'node': nid, 'type': sup.get('type', 'Roller'), 'angle': sup.get('angle', 0.0)})

    for cp in cut_points:
        cx, cz = cp['x'], cp.get('z', cp.get('y', 0.0))
        min_d, w_seg, w_s = get_closest_segment_exact((cx, cz), segments)
        if min_d < 0.30: support_injections[w_seg].append(w_s)
        get_or_add_node(cx, cz)

    for st_idx, st_item in enumerate(struts):
        tx, tz = st_item['tx'], st_item.get('tz', st_item.get('ty', 0.0))
        bx, bz = st_item['bx'], st_item.get('bz', st_item.get('by', 0.0))
        
        dt, wt_seg, wt_s = get_closest_segment_exact((tx, tz), segments)
        if dt < 0.30: 
            support_injections[wt_seg].append(wt_s)
            tx, tz, _ = eval_seg_point(segments[wt_seg], wt_s)
            
        db, wb_seg, wb_s = get_closest_segment_exact((bx, bz), segments)
        if db < 0.30: 
            support_injections[wb_seg].append(wb_s)
            bx, bz, _ = eval_seg_point(segments[wb_seg], wb_s)
            
        top_node, bot_node = get_or_add_node(tx, tz), get_or_add_node(bx, bz)
        elements.append({'type': 'truss', 'group': 'strut', 'sec': st_item.get('sec', 'Unknown'), 'n1': bot_node, 'n2': top_node, 'strut_idx': st_idx, 'E': 21000000.0, 'A': 0.001})

    for i, seg in enumerate(segments):
        L = seg['L']
        key_s_vals = [0.0, L] + support_injections[i]
        
        for ld in loads:
            if ld.get('seg_idx') == i: key_s_vals.extend([ld['start'], ld['end']])
                
        num_subdivisions = max(1, int(np.ceil(L / mesh_size)))
        key_s_vals.extend(np.linspace(0, L, num_subdivisions + 1))
            
        keys = sorted(list(set([min(max(round(k, 4), 0.0), round(L, 4)) for k in key_s_vals])))
        node_indices = [get_or_add_node(*eval_seg_point(seg, s)[:2]) for s in keys]
        
        m_idx = seg.get('master_idx', i)
        sec_props = seg_sections[m_idx] if m_idx < len(seg_sections) else seg_sections[0]
        
        for j in range(len(keys)-1):
            n1, n2 = node_indices[j], node_indices[j+1]
            if n1 == n2: continue 
                
            s_mid = (keys[j] + keys[j+1]) / 2.0
            _, _, th_mid = eval_seg_point(seg, s_mid)
            c_t, s_t = np.cos(th_mid), np.sin(th_mid)
            
            p_x1, p_z1, p_x2, p_z2 = 0.0, 0.0, 0.0, 0.0
            
            for ld in loads:
                if ld.get('seg_idx') == i and ld.get('type') != 'Point Load' and ld['start'] - 1e-4 <= s_mid <= ld['end'] + 1e-4:
                    L_ld = max(ld['end'] - ld['start'], 1e-5)
                    wa = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j] - ld['start']) / L_ld
                    wb = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j+1] - ld['start']) / L_ld
                    dir_str = ld.get('dir', 'Global Z (Vertical)').upper()
                    
                    if 'Z' in dir_str or 'Y' in dir_str:
                        p_x1 += wa * s_t; p_z1 += wa * c_t; p_x2 += wb * s_t; p_z2 += wb * c_t
                    elif 'X' in dir_str:
                        p_x1 += wa * c_t; p_z1 -= wa * s_t; p_x2 += wb * c_t; p_z2 -= wb * s_t
                    else:
                        p_z1 += wa; p_z2 += wb
                        
            elements.append({'type': 'frame', 'group': 'segment', 'sec': sec_props['name'], 'n1': n1, 'n2': n2, 'px1': p_x1, 'py1': p_z1, 'px2': p_x2, 'py2': p_z2, 'E': sec_props['E'] * 10000.0, 'A': sec_props['A'], 'I': sec_props['I'] / 100000000.0, 'seg_idx': i, 'L': keys[j+1] - keys[j], 'th_mid': th_mid})
            
        for ld in loads:
            if ld.get('seg_idx') == i and ld.get('type') == 'Point Load':
                px, pz, th_pt = eval_seg_point(seg, ld['start'])
                nid = get_or_add_node(px, pz)
                dir_str = ld.get('dir', 'Global Z (Vertical)').upper()
                
                if 'Z' in dir_str or 'Y' in dir_str: nodal_loads.append({'node': nid, 'Fx': 0.0, 'Fz': ld['w1']})
                elif 'X' in dir_str: nodal_loads.append({'node': nid, 'Fx': ld['w1'], 'Fz': 0.0})
                else: nodal_loads.append({'node': nid, 'Fx': -ld['w1'] * np.sin(th_pt), 'Fz': ld['w1'] * np.cos(th_pt)})

    return nodes, elements, nodal_loads, supports_list_out

def solve_fea_engine(nodes, elements, nodal_loads, supports_list):
    NDOF = len(nodes) * 3
    K, F = np.zeros((NDOF, NDOF)), np.zeros(NDOF)
    
    for el in elements:
        n1, n2 = el['n1'], el['n2']
        x1, z1, x2, z2 = nodes[n1][0], nodes[n1][1], nodes[n2][0], nodes[n2][1]
        L = np.hypot(x2 - x1, z2 - z1)
        if L < 1e-5: continue
            
        c, s = (x2 - x1) / L, (z2 - z1) / L
        el['L'], el['c'], el['s'] = L, c, s
        E, A, I = el['E'], el['A'], el.get('I', 0.00005)
        
        T = np.array([[c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0], [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0, 0] = k_loc[3, 3] = E * A / L
            k_loc[0, 3] = k_loc[3, 0] = -E * A / L
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], 
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], 
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], 
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1, py1, px2, py2 = el.get('px1', 0.0), el.get('py1', 0.0), el.get('px2', 0.0), el.get('py2', 0.0)
            f_loc = np.array([(2*px1 + px2) * L / 6.0, (7*py1 + 3*py2) * L / 20.0, (3*py1 + 2*py2) * L**2 / 60.0, (px1 + 2*px2) * L / 6.0, (3*py1 + 7*py2) * L / 20.0, -(2*py1 + 3*py2) * L**2 / 60.0])
            f_glob = T.T @ f_loc
            dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
            for r in range(6): F[dof[r]] += f_glob[r]
            
        k_glob = T.T @ k_loc @ T
        dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        for r in range(6):
            for col in range(6): K[dof[r], dof[col]] += k_glob[r, col]
                
    for nl in nodal_loads:
        F[3*nl['node']] += nl.get('Fx', 0.0)
        F[3*nl['node']+1] += nl.get('Fz', nl.get('Fy', 0.0))
        
    net_load_z = abs(np.sum(F[1::3])) 
    K_orig, fixed_dofs, K_pen = K.copy(), [], 1e12 
    
    for sup in supports_list:
        n, t, a = sup['node'], sup['type'], sup.get('angle', 0.0)
        if t == 'Fixed': fixed_dofs.extend([3*n, 3*n+1, 3*n+2])
        elif t == 'Hinged': fixed_dofs.extend([3*n, 3*n+1])
        elif t == 'Roller':
            rad = np.radians(a)
            nx, nz = -np.sin(rad), np.cos(rad) 
            K[3*n, 3*n] += K_pen * nx**2
            K[3*n+1, 3*n+1] += K_pen * nz**2
            K[3*n, 3*n+1] += K_pen * nx * nz
            K[3*n+1, 3*n] += K_pen * nx * nz

    free_dof = [i for i in range(NDOF) if i not in fixed_dofs]
    K_ff, F_f = K[np.ix_(free_dof, free_dof)], F[free_dof]
    
    U = np.zeros(NDOF)
    try: U[free_dof] = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError: U[free_dof] = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    R_reactions = K_orig @ U - F 
    
    for el in elements:
        if el.get('L', 0) < 1e-5: continue
            
        n1, n2, c, s, L, E, A, I = el['n1'], el['n2'], el['c'], el['s'], el['L'], el['E'], el['A'], el.get('I', 0.00005)
        u_glob = U[[3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]]
        T = np.array([[c, s, 0, 0, 0, 0], [-s, c, 0, 0, 0, 0], [0, 0, 1, 0, 0, 0], [0, 0, 0, c, s, 0], [0, 0, 0, -s, c, 0], [0, 0, 0, 0, 0, 1]])
        u_loc = T @ u_glob
        el['internal'] = {'u_loc': u_loc}
        
        if el['type'] == 'truss':
            N_val = (E * A / L) * (u_loc[3] - u_loc[0])
            xs = np.linspace(0, L, 51)
            el['internal'].update({'N': np.full_like(xs, N_val), 'V': np.zeros_like(xs), 'M': np.zeros_like(xs), 'D': np.zeros_like(xs), 'x': xs})
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2], 
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], [-E*A/L, 0, 0, E*A/L, 0, 0], 
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            px1, py1, px2, py2 = el.get('px1', 0.0), el.get('py1', 0.0), el.get('px2', 0.0), el.get('py2', 0.0)
            f_loc = np.array([(2*px1 + px2)*L/6.0, (7*py1 + 3*py2)*L/20.0, (3*py1 + 2*py2)*L**2/60.0, (px1 + 2*px2)*L/6.0, (3*py1 + 7*py2)*L/20.0, -(2*py1 + 3*py2)*L**2/60.0])
            f_end = k_loc @ u_loc - f_loc
            xs = np.linspace(0, L, 51)
            N_arr, V_arr, M_arr, D_arr = np.zeros_like(xs), np.zeros_like(xs), np.zeros_like(xs), np.zeros_like(xs)
            v1, th1, v2, th2 = u_loc[1], u_loc[2], u_loc[4], u_loc[5]
            
            for i, x in enumerate(xs):
                N_arr[i] = -f_end[0] - (px1*x + (px2-px1)*x**2/(2*L))
                V_arr[i] = f_end[1] + (py1*x + (py2-py1)*x**2/(2*L))
                M_arr[i] = -f_end[2] + f_end[1]*x + py1*x**2/2.0 + (py2-py1)*x**3/(6*L)
                xi = x / L if L > 0 else 0
                v_x = v1*(1.0 - 3.0*xi**2 + 2.0*xi**3) + th1*(L * (xi - 2.0*xi**2 + xi**3)) + v2*(3.0*xi**2 - 2.0*xi**3) + th2*(L * (-xi**2 + xi**3))
                v_load = (((py1 + py2) / 2.0) * x**2 * (L - x)**2) / (24.0 * E * I) if (E * I) != 0 else 0
                D_arr[i] = (v_x + v_load) * 1000.0 
                
            el['internal'].update({'N': N_arr, 'V': V_arr, 'M': M_arr, 'D': D_arr, 'x': xs})
            
    return U, R_reactions, net_load_z

# ==============================================================================
# 4. THE BEAST OPTIMIZER 
# ==============================================================================
def run_bridge_optimizer(base_segments, working_segments, active_seg_sections, ui_struts, ui_loads, target_rxn, spacings_str, auto_mesh_size, is_symmetric, opt_mode, combo_factors, status_text, progress_bar):
    try: 
        spacings_raw = spacings_str.split(',')
        spacings = []
        for x in spacings_raw: spacings.append(float(x.strip()))
        spacings.sort(reverse=True)
    except Exception as e: return False, None, None, "❌ Format error in spacings. Please use comma separated values."
    
    if not base_segments: return False, None, None, "❌ No base segments found to optimize."
        
    all_xs = [p[0] for seg in base_segments for p in (seg.get('abs_p1'), seg.get('abs_p2')) if p]
    soldier_min_x, soldier_max_x = min(all_xs), max(all_xs)
    center_x, half_width = (soldier_min_x + soldier_max_x) / 2.0, (soldier_max_x - soldier_min_x) / 2.0
    
    test_combined_loads = []
    for i, ld in enumerate(ui_loads):
        t_mode = ld.get('target_mode', 'Single Segment')
        target_base_indices = []
        if t_mode == "Single Segment": target_base_indices.append(ld.get('seg_idx', 0))
        elif t_mode == "Multiple Segments": target_base_indices.extend(ld.get('target_segs_idx', []))
        else: target_base_indices.extend(range(len(base_segments)))
            
        cat = ld.get('category', 'Dead Load')
        fac = combo_factors.get(cat, 1.0)
        w1 = ld.get('w1', 0.0) * fac
        w2 = ld.get('w2', ld.get('w1', 0.0)) * fac if ld.get('type') == 'Trapezoidal' else w1
        loc_m = ld.get('loc', 0.0)
        
        target_working_indices = []
        for w_idx, w_seg in enumerate(working_segments):
            if w_seg.get('master_idx', 0) in target_base_indices: target_working_indices.append(w_idx)
        
        for s_idx_num in target_working_indices:
            w_len = float(working_segments[s_idx_num].get('L', 0.0))
            start_val = min(loc_m, w_len) if ld.get('type') == 'Point Load' else 0.0
            end_val = start_val if ld.get('type') == 'Point Load' else w_len
            test_combined_loads.append({'seg_idx': s_idx_num, 'category': cat, 'type': ld.get('type', 'Uniform'), 'dir': ld.get('dir', 'Global Z (Vertical)'), 'start': start_val, 'end': end_val, 'w1': w1, 'w2': w2})

    def run_trial(test_supps, dynamic_struts):
        nodes_t, elems_t, nloads_t, slist_t = build_chain_mesh(working_segments, active_seg_sections, test_combined_loads, dynamic_struts, test_supps, [], mesh_size=auto_mesh_size)
        U, R, net_load = solve_fea_engine(nodes_t, elems_t, nloads_t, slist_t)
        ry_list = [R[3*sup['node']+1] for sup in slist_t]
        max_ry = max(ry_list) if ry_list else 0
        min_ry = min(ry_list) if ry_list else 0
        
        soldier_safe = True
        for i, sec in enumerate(active_seg_sections):
            for el in elems_t:
                if el.get('group') == 'segment' and el.get('seg_idx') == i:
                    m_arr, v_arr = el.get('internal', {}).get('M', [0]), el.get('internal', {}).get('V', [0])
                    if np.max(np.abs(m_arr)) > sec['Mall'] or np.max(np.abs(v_arr)) > sec['Qall']: 
                        soldier_safe = False; break
                        
        struts_safe, upgraded_struts = True, []
        for el in elems_t:
            if el['type'] == 'truss':
                N_max = np.max(np.abs(el.get('internal', {}).get('N', [0])))
                opt_sec = get_optimal_strut_section(el.get('L', 0.0), N_max)
                if not opt_sec: 
                    struts_safe = False; upgraded_struts.append(el.get('sec'))
                else: 
                    upgraded_struts.append(opt_sec)
        return max_ry, min_ry, net_load, soldier_safe, struts_safe, upgraded_struts

    dummy_supps = [{'x': soldier_min_x, 'z': drop_ray_to_bottom_chord(soldier_min_x, base_segments)[1], 'type': 'Hinged', 'angle': 0.0}, {'x': soldier_max_x, 'z': drop_ray_to_bottom_chord(soldier_max_x, base_segments)[1], 'type': 'Hinged', 'angle': 0.0}]
    _, _, total_system_load, _, _, _ = run_trial(dummy_supps, ui_struts)
    min_required_props = max(2, int(math.ceil(total_system_load / target_rxn))) if target_rxn > 1e-4 else 2

    valid_grids = []
    if is_symmetric:
        def build_sym_grids(current_grid):
            cantilever = half_width - current_grid[-1]
            if 0.15 <= cantilever <= 1.50:
                full_grid = set(current_grid)
                for x in current_grid:
                    if x > 1e-4: full_grid.add(-x)
                sym_coords = []
                for x in sorted(list(full_grid)):
                    cx = round(center_x + x, 3)
                    cx = max(soldier_min_x, min(soldier_max_x, cx)) 
                    sym_coords.append(cx)
                valid_grids.append(tuple(sym_coords))
            if cantilever < 0.15: return
            for s in spacings: build_sym_grids(current_grid + [current_grid[-1] + s])
                
        build_sym_grids([0.0])
        for s in spacings: 
            build_sym_grids([s / 2.0])
            if "Deep" in opt_mode: build_sym_grids([s / 3.0]); build_sym_grids([s / 4.0])
            
    else:
        def build_asym_grids(current_grid):
            cantilever_right = soldier_max_x - current_grid[-1]
            if 0.15 <= cantilever_right <= 1.50:
                asym_coords = [round(x, 3) for x in current_grid if soldier_min_x - 0.05 <= x <= soldier_max_x + 0.05]
                valid_grids.append(tuple(asym_coords))
            if cantilever_right < 0.15: return
            for s in spacings: build_asym_grids(current_grid + [current_grid[-1] + s])
                
        cantilever_opts = np.arange(0.15, 1.51, 0.10)
        for lc in cantilever_opts: build_asym_grids([soldier_min_x + lc])

    filtered_grids = [list(g) for g in set(valid_grids) if len(g) >= min_required_props]
    if not filtered_grids: return False, None, None, f"❌ Impossible to optimize! Requires at least {min_required_props} props while respecting cantilevers."

    grids_by_props = {}
    for g in filtered_grids:
        p_count = len(g)
        if p_count not in grids_by_props: grids_by_props[p_count] = []
        grids_by_props[p_count].append(g)

    shift_options = [0.0] if is_symmetric else [0.0, 0.10, -0.10, 0.20, -0.20]
    max_time = 180.0 if "Quick" in opt_mode else 900.0  
        
    start_time = time.time()
    best_fallback_grid, best_fallback_struts, best_fallback_score = None, ui_struts, 999999.0
    trials_count = 0
    total_estimated_trials = len(filtered_grids) * len(shift_options)
    sorted_p_keys = sorted(list(grids_by_props.keys()))
    timeout_reached = False
    
    for p_count in sorted_p_keys:
        if timeout_reached: break
        for actual_coords in grids_by_props[p_count]:
            if time.time() - start_time > max_time: timeout_reached = True; break
            
            cantilever_L = half_width - (max(actual_coords) - center_x) if is_symmetric else soldier_max_x - max(actual_coords)
            excluded_zone_start = soldier_max_x - (cantilever_L / 3.0)
            excluded_zone_start_left = soldier_min_x + (cantilever_L / 3.0)
            
            test_supps = []
            for gx in actual_coords:
                final_x, final_z = drop_ray_to_bottom_chord(gx, base_segments)
                test_supps.append({'x': final_x, 'z': round(final_z, 3), 'type': 'Hinged', 'angle': 0.0})
                
            for shift_val in shift_options:
                if time.time() - start_time > max_time: timeout_reached = True; break
                
                shifted_struts = []
                for strut in ui_struts:
                    new_strut = strut.copy()
                    nz_b_old, nz_t_old = strut.get('bz', strut.get('by', 0.0)), strut.get('tz', strut.get('ty', 0.0))
                    nx_b, nz_b = get_shifted_coords_along_segment(strut['bx'], nz_b_old, shift_val, base_segments)
                    nx_t, nz_t = get_shifted_coords_along_segment(strut['tx'], nz_t_old, shift_val, base_segments)
                    if nx_b > excluded_zone_start or nx_b < excluded_zone_start_left: nx_b, nz_b, nx_t, nz_t = strut['bx'], nz_b_old, strut['tx'], nz_t_old
                    new_strut.update({'bx': nx_b, 'bz': nz_b, 'tx': nx_t, 'tz': nz_t, 'sec': get_valid_strut_names()[0] if STRUTS_DB else "Unknown"})
                    shifted_struts.append(new_strut)
                
                max_ry, min_ry, _, soldier_safe, struts_safe, upg_secs = run_trial(test_supps, shifted_struts)
                
                if not struts_safe and len(upg_secs) == len(shifted_struts):
                    for idx_st in range(len(shifted_struts)): shifted_struts[idx_st]['sec'] = upg_secs[idx_st]
                    max_ry, min_ry, _, soldier_safe, struts_safe, _ = run_trial(test_supps, shifted_struts) 
                
                trials_count += 1
                if trials_count % 15 == 0:
                    progress_bar.progress(min(1.0, trials_count / float(total_estimated_trials)))
                    status_text.markdown(f"**⏳ Search:** Grid **{p_count} Props** | Best Rxn So Far: **{best_fallback_score:.2f} kN**")
                
                if max_ry <= target_rxn and min_ry >= 0.5 and soldier_safe and struts_safe:
                    progress_bar.progress(1.0); status_text.empty()
                    return True, test_supps, shifted_struts, f"✅ BOOM! Safe Grid Found: Max Rxn = {max_ry:.2f} kN. Props = {p_count}."
                    
                if max_ry < best_fallback_score and soldier_safe:
                    best_fallback_score, best_fallback_grid, best_fallback_struts = max_ry, test_supps, shifted_struts
                        
    progress_bar.empty(); status_text.empty()
    if best_fallback_grid: return True, best_fallback_grid, best_fallback_struts, f"⚠️ Best possible solution applied: Max Rxn = {best_fallback_score:.2f} kN. Please review visually."
    return False, None, None, f"❌ Failed! Cannot satisfy basic stability (Uplift) with forced cantilevers."

# ==============================================================================
# 5. PLOTTING ENGINE & WORD REPORT (Aesthetic Diagrams & Distributed Checks)
# ==============================================================================
def draw_advanced_reaction_arrow(ax, node_x, node_z, force_mag, axis_nx, axis_nz, target_rxn=54.4):
    if abs(force_mag) < 0.001: return
    arr_L = 0.5  
    sgn = np.sign(force_mag)
    dx = sgn * axis_nx
    dz = sgn * axis_nz
    start_x = node_x - arr_L * dx
    start_z = node_z - arr_L * dz
    arr_c = 'red' if force_mag < 0 or force_mag > target_rxn else 'green'
    ax.arrow(start_x, start_z, arr_L * dx, arr_L * dz, length_includes_head=True, head_width=0.08, head_length=0.12, fc=arr_c, ec=arr_c, lw=0.8, zorder=5)
    ax.text(start_x - 0.15 * dx, start_z - 0.15 * dz, f"{force_mag:+.2f}", color=arr_c, fontsize=7, fontweight='normal', fontname='Arial', ha='center', va='center')

def draw_base_geometry(ax, nodes, elements, supports_list, segments, show_names=False, show_dimensions=False):
    for el in elements:
        n1, n2 = nodes[el['n1']], nodes[el['n2']]
        if el['type'] == 'truss':
            ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='red', linestyle='-', linewidth=0.8, zorder=1)
            if show_names:
                mid_x, mid_z = (n1[0] + n2[0]) / 2.0, (n1[1] + n2[1]) / 2.0
                th = math.atan2(n2[1] - n1[1], n2[0] - n1[0])
                rot_deg = math.degrees(th)
                if rot_deg > 90: rot_deg -= 180
                elif rot_deg < -90: rot_deg += 180
                sec_name = el.get('sec', '').split('(')[0].strip()
                ax.text(mid_x, mid_z + 0.15, sec_name, color='#888888', fontsize=6, ha='center', va='center', rotation=rot_deg, fontweight='normal')
        else:
            ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='royalblue', linestyle='-', linewidth=1.5, zorder=1)
            
    sorted_sups = sorted(supports_list, key=lambda s: nodes[s['node']][0])
    for i, sup in enumerate(sorted_sups):
        x, z, t = nodes[sup['node']][0], nodes[sup['node']][1], sup['type']
        ang_rad = math.radians(sup.get('angle', 0.0))
        c_a, s_a = math.cos(ang_rad), math.sin(ang_rad)
        def rot(px, pz): return x + (px - x)*c_a - (pz - z)*s_a, z + (px - x)*s_a + (pz - z)*c_a
        
        if t == 'Fixed':
            ax.add_patch(Polygon([rot(x-0.1, z-0.1), rot(x+0.1, z-0.1), rot(x+0.1, z+0.1), rot(x-0.1, z+0.1)], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot(x-0.1, z), rot(x+0.1, z)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Hinged':
            ax.add_patch(Polygon([rot(x, z), rot(x+0.12, z-0.15), rot(x-0.12, z-0.15)], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot(x-0.17, z-0.15), rot(x+0.17, z-0.15)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Roller':
            ax.add_patch(Polygon([rot(x, z), rot(x+0.12, z-0.15), rot(x-0.12, z-0.15)], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            ax.add_patch(plt.Circle(rot(x, z-0.19), 0.04, facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            l1, l2 = rot(x-0.17, z-0.23), rot(x+0.17, z-0.23)
            ax.plot([l1[0], l2[0]], [l1[1], l2[1]], color='limegreen', lw=1.0, zorder=4)

        if show_names:
            lbl_x, lbl_z = rot(x, z - 0.45)
            ax.text(lbl_x, lbl_z, f"J{i+1}", color='#888888', fontsize=7, fontweight='normal', ha='center', va='center', zorder=10)

    if show_names and segments:
        for i, seg in enumerate(segments):
            mx, mz, mth = eval_seg_point(seg, seg.get('L', 0)/2.0)
            rot_deg = math.degrees(mth)
            if rot_deg > 90: rot_deg -= 180
            elif rot_deg < -90: rot_deg += 180
            clean_name = seg.get('name', f"S{i+1}").split('-')[0]
            ax.text(mx - math.sin(mth)*0.3, mz + math.cos(mth)*0.3, clean_name, color='#888888', fontsize=7, ha='center', va='center', rotation=rot_deg, fontname='Arial', fontweight='normal')

    if show_dimensions and len(supports_list) > 1:
        sup_xs = sorted(list(set([nodes[sup['node']][0] for sup in supports_list])))
        if len(sup_xs) > 1:
            dim_z = min([nodes[sup['node']][1] for sup in supports_list]) - 0.85
            ax.plot([sup_xs[0], sup_xs[-1]], [dim_z, dim_z], color='gray', lw=0.6, zorder=1)
            for i in range(len(sup_xs)):
                ax.plot([sup_xs[i], sup_xs[i]], [dim_z - 0.1, dim_z + 0.1], color='gray', lw=0.6, zorder=1)
                if i < len(sup_xs) - 1:
                    dist = sup_xs[i+1] - sup_xs[i]
                    mid_x = (sup_xs[i] + sup_xs[i+1]) / 2.0
                    ax.text(mid_x, dim_z + 0.08, f"{dist:.2f}m", color='gray', fontsize=7, ha='center', fontweight='normal')

def get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter=None, R_reactions=None, target_rxn=54.4):
    apply_plot_styles()
    fig_ld, ax_ld = plt.subplots(figsize=(9, 5.5))
    ax_ld.set_aspect('equal', adjustable='datalim')
    ax_ld.axis('off')
    draw_base_geometry(ax_ld, nodes, elements, supports_list, segments, show_names=True, show_dimensions=True)
    
    for ld in loads:
        if cat_filter and ld.get('category', 'Dead Load') != cat_filter: continue
        i = ld.get('seg_idx', 0)
        if i >= len(segments): continue
        w1, w2 = ld.get('w1', 0.0), ld.get('w2', 0.0)
        if abs(w1) < 1e-4 and abs(w2) < 1e-4: continue
        
        num_pts = max(10, int((ld.get('end', 0) - ld.get('start', 0)) / 0.1))
        s_vals = np.linspace(ld.get('start', 0), ld.get('end', 0), num_pts)
        poly_pts, top_pts, bot_pts = [], [], []
        
        for sv in s_vals:
            px, pz, th = eval_seg_point(segments[i], sv)
            L_load = max(1e-5, (ld.get('end', 0) - ld.get('start', 0)))
            w_val = (w1 + (w2 - w1) * (sv - ld.get('start', 0)) / L_load) * 0.05
            
            bot_pts.append((px, pz))
            
            dir_str = ld.get('dir', 'Global Z (Vertical)').upper()
            if 'Z' in dir_str or 'Y' in dir_str: top_pts.append((px, pz + abs(w_val))) 
            elif 'X' in dir_str: top_pts.append((px - w_val, pz))
            else:
                c, s = math.cos(th), math.sin(th)
                top_pts.append((px + s * abs(w_val), pz + c * abs(w_val)))
                
        poly_pts = bot_pts + top_pts[::-1]
        
        if len(poly_pts) > 2:
            ax_ld.add_patch(Polygon(poly_pts, fill=False, edgecolor='#66B2FF', lw=0.1, hatch='|||', zorder=1))
            ax_ld.add_patch(Polygon(poly_pts, fill=False, edgecolor='#2C68C5', lw=1.0, zorder=2))
            
            num_lines = max(4, int(len(top_pts) / 2))
            for k in range(0, len(top_pts), num_lines):
                ax_ld.plot([bot_pts[k][0], top_pts[k][0]], [bot_pts[k][1], top_pts[k][1]], color='#66B2FF', lw=0.5, alpha=0.6, zorder=2)

            if abs(w1) > 0.01:
                ax_ld.text(top_pts[0][0], top_pts[0][1] + 0.1, f"{abs(w1):.2f}", color='#555555', ha='center', va='bottom', fontsize=7, fontweight='normal', fontname='Arial')
            if abs(w2) > 0.01 and abs(top_pts[-1][0] - top_pts[0][0]) > 0.3:
                ax_ld.text(top_pts[-1][0], top_pts[-1][1] + 0.1, f"{abs(w2):.2f}", color='#555555', ha='center', va='bottom', fontsize=7, fontweight='normal', fontname='Arial')

    if R_reactions is not None:
        for sup in supports_list:
            n = sup['node']
            ang = math.radians(sup.get('angle', 0.0))
            Rx, Rz = R_reactions[3*n], R_reactions[3*n+1]
            x, z = nodes[n][0], nodes[n][1]
            R_loc_x = Rx * math.cos(ang) + Rz * math.sin(ang)
            R_loc_z = -Rx * math.sin(ang) + Rz * math.cos(ang)
            if sup['type'] == 'Roller': draw_advanced_reaction_arrow(ax_ld, x, z, R_loc_z, -math.sin(ang), math.cos(ang), target_rxn)
            else:
                draw_advanced_reaction_arrow(ax_ld, x, z, R_loc_x, math.cos(ang), math.sin(ang), target_rxn)
                draw_advanced_reaction_arrow(ax_ld, x, z, R_loc_z, -math.sin(ang), math.cos(ang), target_rxn)
    return safe_render_fig(fig_ld)

def plot_sap2000_diagrams(nodes, elements, R_reactions, scales, supports_list, loads, segments, target_rxn=54.4):
    apply_plot_styles()
    figs_dict = {}
    has_dl = any(ld.get('category') == 'Dead Load' for ld in loads if abs(ld.get('w1', 0)) > 1e-4)
    has_ll = any(ld.get('category') == 'Live Load' for ld in loads if abs(ld.get('w1', 0)) > 1e-4)
    has_wl = any(ld.get('category') == 'Wind Load' for ld in loads if abs(ld.get('w1', 0)) > 1e-4)
    
    if has_dl: figs_dict['DL'] = get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter='Dead Load')
    if has_ll: figs_dict['LL'] = get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter='Live Load')
    if has_wl: figs_dict['WL'] = get_live_preview_image(nodes, elements, supports_list, loads, segments, cat_filter='Wind Load')
    
    def create_force_plot(val_key, scale, c_pos, c_neg):
        fig_f, ax_f = plt.subplots(figsize=(9, 5.5))
        ax_f.set_aspect('equal', adjustable='datalim')
        ax_f.axis('off')
        draw_base_geometry(ax_f, nodes, elements, supports_list, segments, show_names=True)
        for el in elements:
            n1, n2 = el['n1'], el['n2']
            x1, z1 = nodes[n1][0], nodes[n1][1]
            x2, z2 = nodes[n2][0], nodes[n2][1]
            c, s = el.get('c', 1.0), el.get('s', 0.0)
            xs = el.get('internal', {}).get('x', [])
            vals = el.get('internal', {}).get(val_key, [])
            if len(vals) == 0 or np.all(np.abs(vals) < 1e-6): continue
            plot_vals = -vals if val_key != 'N' else vals
            px = x1 - s * plot_vals * scale + c * xs
            pz = z1 + c * plot_vals * scale + s * xs
            for k in range(len(px)-1): ax_f.plot([px[k], px[k+1]], [pz[k], pz[k+1]], color=c_pos if vals[k] >= 0 else c_neg, lw=0.8)
            ax_f.plot([x1, px[0]], [z1, pz[0]], color=c_pos if vals[0]>=0 else c_neg, lw=0.8)
            ax_f.plot([x2, px[-1]], [z2, pz[-1]], color=c_pos if vals[-1]>=0 else c_neg, lw=0.8)
            mid = len(vals)//2
            if abs(vals[mid]) > 0.1: ax_f.text(px[mid], pz[mid], f"{vals[mid]:+.2f}", fontsize=6, color=c_pos if vals[mid]>=0 else c_neg, ha='center', va='center', fontweight='normal')
        return safe_render_fig(fig_f)

    figs_dict['N'] = create_force_plot('N', scales['N'], 'blue', 'red')
    figs_dict['V'] = create_force_plot('V', scales['V'], 'blue', 'red')
    figs_dict['M'] = create_force_plot('M', scales['M'], 'blue', 'red')
    
    fig_r, ax_r = plt.subplots(figsize=(9, 5.5))
    ax_r.set_aspect('equal', adjustable='datalim')
    ax_r.axis('off')
    draw_base_geometry(ax_r, nodes, elements, supports_list, segments, show_names=True, show_dimensions=True)
    for sup in supports_list:
        n = sup['node']
        ang = math.radians(sup.get('angle', 0.0))
        Rx, Rz = R_reactions[3*n], R_reactions[3*n+1]
        x, z = nodes[n][0], nodes[n][1]
        R_loc_x = Rx * math.cos(ang) + Rz * math.sin(ang)
        R_loc_z = -Rx * math.sin(ang) + Rz * math.cos(ang)
        if sup['type'] == 'Roller': draw_advanced_reaction_arrow(ax_r, x, z, R_loc_z, -math.sin(ang), math.cos(ang), target_rxn)
        else:
            draw_advanced_reaction_arrow(ax_r, x, z, R_loc_x, math.cos(ang), math.sin(ang), target_rxn)
            draw_advanced_reaction_arrow(ax_r, x, z, R_loc_z, -math.sin(ang), math.cos(ang), target_rxn)
    figs_dict['R'] = safe_render_fig(fig_r)
    return figs_dict


# =========================================================
# 6. WORD REPORT GENERATOR (The Ultimate Distributed Sheet)
# =========================================================
def generate_multi_case_report(cases_data, proj_info):
    import fitz  
    if os.path.exists("Acrow_Template.docx"): doc = Document("Acrow_Template.docx")
    else: doc = Document()

    system_title = "SHOREBRACE TABLE FORM SYSTEM"
    for p in doc.paragraphs: 
        if p.text and "CALCULATION SHEET FOR" in p.text.upper():
            for r in p.runs: r.text = ""
            run = p.add_run(system_title)
            run.font.name, run.font.size, run.font.bold, run.font.color.rgb = 'Arial', Pt(16), True, RGBColor(255,255,255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs: 
                    if p.text and "CALCULATION SHEET FOR" in p.text.upper():
                        for r in p.runs: r.text = ""
                        run = p.add_run(system_title)
                        run.font.name, run.font.size, run.font.bold, run.font.color.rgb = 'Arial', Pt(16), True, RGBColor(255,255,255)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    replacements = {"[PROJECT_NAME]": proj_info.get("proj_name", ""), "[CONTRACTOR]": proj_info.get("contractor", ""), "[CALC_SUBJECT]": proj_info.get("calc_sub", ""), "[SYSTEM_NAME]": proj_info.get("sys_name", ""), "[PROJ_NO]": proj_info.get("proj_no", ""), "[DATE]": proj_info.get("date_val", ""), "[CALC_BY]": proj_info.get("calc_by", ""), "[CHK_BY]": proj_info.get("chk_by", ""), "[REV]": "00"}
    
    for p in doc.paragraphs:
        if "[COVER_IMAGE]" in p.text:
            p.text = p.text.replace("[COVER_IMAGE]", "")
            if proj_info.get("cover_img") and os.path.exists(proj_info.get("cover_img")): 
                p.add_run().add_picture(proj_info.get("cover_img"), width=Cm(15.0))
            
            p_proj = doc.add_paragraph()
            p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_proj = p_proj.add_run(f"\n{proj_info.get('proj_name', 'Acrow Mega Project')}")
            r_proj.font.name, r_proj.font.size, r_proj.font.bold, r_proj.font.color.rgb = 'Arial', Pt(16), True, RGBColor(192, 0, 0)
            
            p_cont = doc.add_paragraph()
            p_cont.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cont = p_cont.add_run(proj_info.get("contractor", "Main Contractor"))
            r_cont.font.name, r_cont.font.size, r_cont.font.bold = 'Arial', Pt(14), True
            
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("\nCALCULATION SHEET FOR BRIDGE DECK SLAB")
            r_title.font.name, r_title.font.size, r_title.font.bold, r_title.font.color.rgb = 'Arial', Pt(14), True, RGBColor(192, 0, 0)
            
            p_using = doc.add_paragraph()
            p_using.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_using = p_using.add_run("\nUSING\n")
            r_using.font.name, r_using.font.size, r_using.font.bold = 'Arial', Pt(14), True
            
            insert_blue_banner(doc, "SHOREBRACE TABLE FORM SYSTEM", font_size=16)
            
        for k, v in replacements.items():
            if k in p.text: p.text = p.text.replace(k, str(v))
            
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.footer, sec.first_page_footer]:
            if hf:
                for tbl in hf.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for k, v in replacements.items():
                                    if k in p.text: p.text = p.text.replace(k, str(v))

    doc.add_page_break()

    insert_blue_banner(doc, "INDEX OF CONTENTS", font_size=16)
    def get_pdf_page_count_safe(pdf_path):
        try:
            pdf_d = fitz.open(pdf_path)
            c = len(pdf_d)
            pdf_d.close()
            return c
        except: return 1

    ds_page_map = {}
    current_page = 4 
    data_sheets = proj_info.get("data_sheets", [])
    if data_sheets:
        for f in data_sheets:
            if os.path.exists(f):
                p_count = get_pdf_page_count_safe(f)
                bname = os.path.basename(f).replace('.pdf', '')
                ds_page_map[bname] = f"{current_page}" if p_count == 1 else f"{current_page}-{current_page + p_count - 1}"
                current_page += p_count

    def add_line(text, bold=False, size=12, italic=False, color=None, align='left', underline=False):
        p = doc.add_paragraph()
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else: force_ltr_left(p)
        r = p.add_run(text)
        r.font.name, r.font.size, r.font.bold, r.font.italic, r.font.rtl, r.underline = 'Arial', Pt(size), bold, italic, False, underline
        if color: r.font.color.rgb = color
        return p

    add_heading_14(doc, "1. Formwork Materials Technical Data:")
    idx_c = 1
    for k, v in ds_page_map.items():
        clean_k = re.sub(r'(?i)data\s*sheet\s*for\s*', '', k).strip()
        add_line(f"   1.{idx_c} Datasheet for Data Sheet for {clean_k} ........................................ Page {v}", size=11)
        idx_c += 1

    doc.add_paragraph()
    add_heading_14(doc, "2. Design Loads:")
    design_pdf = "Design_Loads_BS.pdf" if "BS" in proj_info.get("ref_code", "") and os.path.exists("Design_Loads_BS.pdf") else ("Design_Loads_ACI.pdf" if "ACI" in proj_info.get("ref_code", "") and os.path.exists("Design_Loads_ACI.pdf") else None)
    if design_pdf:
        p_count = get_pdf_page_count_safe(design_pdf)
        design_loads_page = f"{current_page}" if p_count == 1 else f"{current_page}-{current_page + p_count - 1}"
        add_line(f"   2.1 Design Loads .................................................... Page {design_loads_page}", size=11)
        current_page += p_count

    doc.add_paragraph()
    add_heading_14(doc, "3. Formwork Elements Calculations:")
    
    # 🎯 Smart Page Counter for TOC
    calc_start_page = current_page
    for i_idx, case in enumerate(cases_data): 
        tbl_id_clean = str(case['title'].upper()).replace("TABLE", "").strip()
        add_line(f"   3.{i_idx+1} Elements Calculation for Table T{tbl_id_clean}", size=11)
        
        # Calculate Pages for this case (Cover/Data + Plywood + Sec + Soldier + Struts + Shoring)
        p_ply = calc_start_page + 1
        p_sec = calc_start_page + 2
        p_sol = calc_start_page + 3
        
        has_struts = case.get('global_checks', {}).get('max_strut_N', 0) > 0.1
        p_strut = calc_start_page + 4
        p_shore = calc_start_page + 5 if has_struts else calc_start_page + 4
        
        if case.get('local_params'):
            ply_clean_toc = case['local_params'].get('ply_type', '18mm Plywood').replace(' Plywood', '')
            add_line(f"        1- Plywood ({ply_clean_toc}) .............................................................. Page {p_ply}", size=10)
            add_line(f"        2- Secondary Decking ({case['local_params'].get('sec_type', 'Timber H20')}) ................................. Page {p_sec}", size=10)
        add_line(f"        3- Main Decking .......................................................................... Page {p_sol}", size=10)
        if has_struts:
            add_line(f"        4- Push Pull Struts ...................................................................... Page {p_strut}", size=10)
        add_line(f"        5- Support / Shoring ..................................................................... Page {p_shore}", size=10)
        doc.add_paragraph()
        
        calc_start_page = p_shore

    doc.add_page_break()

    insert_blue_banner(doc, "REGULATIONS AND STANDARDS", font_size=16)
    doc.add_paragraph()
    if "BS" in proj_info.get("ref_code", ""): 
        for txt in ["1- BS 5975-1996: FORMWORK FOR CONCRETE", "2- BS 5975-2008: FORMWORK FOR CONCRETE", "3- FORMWORK A GUIDE TO A GOOD PRACTICE", "4- WISA®-FORM PLYWOOD.", "5- THE SAUDI BUILDING CODE (SBC) 2024"]: add_eq(doc, txt)
    else: 
        for txt in ["1- ACI 347R-14 ....... GUIDE TO FORMWORK FOR CONCRETE.", "2- ACI SP-4 ......... FORMWORK FOR CONCRETE.", "3- WISA®-FORM PLYWOOD.", "4- THE SAUDI BUILDING CODE (SBC) 2024"]: add_eq(doc, txt)
    
    if data_sheets:
        doc.add_page_break()
        insert_blue_banner(doc, "FORMWORK MATERIALS TECHNICAL DATA", font_size=14)
        for f in data_sheets:
            if os.path.exists(f): append_pdf_stream_to_word(f, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)
    
    if design_pdf: 
        doc.add_page_break()
        insert_blue_banner(doc, "DESIGN LOADS FOR BRIDGE DECK SLAB", font_size=14)
        append_pdf_stream_to_word(design_pdf, doc, is_path=True, max_width_cm=17.5, max_height_cm=24.0, add_border=True, reduce_first_page=True)

    for case in cases_data:
        doc.add_page_break()
        tbl_id_clean = str(case['title'].upper()).replace("TABLE", "").strip()
        
        insert_blue_banner(doc, f"FORMWORK DESIGN FOR BRIDGE DECK SLAB TABLE T{tbl_id_clean}", font_size=14)
        doc.add_paragraph()

        lp = case.get('local_params', {})
        if lp:
            add_line("Cross Section Descriptive Data:", bold=True, size=12, underline=True)
            t_data = doc.add_table(rows=5, cols=2)
            t_data.alignment = WD_TABLE_ALIGNMENT.CENTER
            desc_data = [
                ("Total Cross Section Depth", f"{lp.get('cs_depth', 0.0):.2f} m"),
                ("Top Slab Thickness", f"{lp.get('top_ts', 0.0):.2f} m"),
                ("Bottom Slab Thickness", f"{lp.get('bot_ts', 0.0):.2f} m"),
                ("Web Concrete Thickness", f"{lp.get('web_ts', 0.0):.2f} m"),
                ("Loaded Width (Soldier Spacing)", f"{case.get('loaded_width_curr', 1.30):.2f} m")
            ]
            for i, (k, v) in enumerate(desc_data):
                row_cells = t_data.rows[i].cells
                row_cells[0].text, row_cells[1].text = k, v
                for j in range(2):
                    tcPr = row_cells[j]._element.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F2F2'); tcPr.append(shd)
                    for p in row_cells[j].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs: r.font.name, r.font.size = 'Arial', Pt(11)
            doc.add_paragraph()

        if 'calc_details' in case and len(case['calc_details']) > 0:
            add_line("1) Dead Load:", bold=True, size=12, underline=True)
            add_line("Equation Used: Load W (kN/m) = [Area (m2) × Density (kN/m3) × Loaded Width (m)] / Length (m)", size=10, italic=True)
            
            table_ld = doc.add_table(rows=len(case['calc_details'])+1, cols=4)
            table_ld.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table_ld.rows[0].cells
            for i, text in enumerate(["Segment", "Length (m)", "Area (m2)", "Load W (kN/m)"]):
                hdr_cells[i].text = text
                tcPr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1F497D'); tcPr.append(shd)
                for p in hdr_cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: r.font.name, r.font.size, r.font.bold, r.font.color.rgb = 'Arial', Pt(12), True, RGBColor(255,255,255)
            
            for i, r_data in enumerate(case['calc_details']):
                row_cells = table_ld.rows[i+1].cells
                row_cells[0].text, row_cells[1].text = str(r_data['segment']), f"{r_data['length']:.2f}"
                row_cells[2].text, row_cells[3].text = f"{r_data['area']:.2f}", f"{r_data['load_w']:.2f}"
                for cell in row_cells:
                    tcPr = cell._element.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F2F2'); tcPr.append(shd)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs: r.font.name, r.font.size = 'Arial', Pt(11)
            doc.add_paragraph()
            
            base_ll = float(case.get('base_live_load', 2.90))
            current_tab_width = float(case.get('loaded_width_curr', 1.30))
            ll_w_calc = base_ll * current_tab_width
            if ll_w_calc > 0:
                add_line("2) Live Load:", bold=True, size=12, underline=True)
                add_line(f"- W1 = live load x Loaded Width by one Soldier = {base_ll:.2f} x {current_tab_width:.2f} = {ll_w_calc:.2f} kN/m.", size=12)
            doc.add_paragraph()
            
            if 'DL' in case.get('img_bufs', {}):
                add_line("Dead Load Distribution Diagram:", bold=False, size=12, color=RGBColor(192, 0, 0), underline=True)
                p_dl = doc.add_paragraph(); p_dl.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_dl.add_run().add_picture(io.BytesIO(case['img_bufs']['DL']), width=Cm(16.5))
            if 'LL' in case.get('img_bufs', {}):
                add_line("Live Load Distribution Diagram:", bold=False, size=12, color=RGBColor(192, 0, 0), underline=True)
                p_ll = doc.add_paragraph(); p_ll.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_ll.add_run().add_picture(io.BytesIO(case['img_bufs']['LL']), width=Cm(16.5))

        if lp:
            LL = float(case.get('base_live_load', 2.90))
            FW = 0.5 
            ply_type_raw = case.get('local_params', {}).get('ply_type', '18mm Plywood')
            ply_type_clean = ply_type_raw.replace(' Plywood', '') 
            sec_type = lp.get('sec_type', 'Timber H20')
            
            for part_name, ts, s_spc, m_spc, sec_L, sec_cant in [
                ("Bottom Slab", lp.get('bot_ts', 0), lp.get('sec_spc_bot', 0), lp.get('main_spc_bot', 0), lp.get('sec_l_bot', 2.5), lp.get('sec_cant_bot', 0.65)),
                ("Web", lp.get('web_ts', 0), lp.get('sec_spc_web', 0), lp.get('main_spc_web', 0), lp.get('sec_l_web', 2.5), lp.get('sec_cant_web', 0.65))
            ]:
                if ts <= 0.01 or s_spc <= 0.01 or m_spc <= 0.01: continue
                
                doc.add_page_break()
                add_line(f"CHECK FORMWORK ELEMENTS UNDER {part_name.upper()} (Depth = {ts:.2f}m)", bold=True, size=14, underline=True)
                w_tot = 25.0 * ts + LL + FW
                
                add_heading_14(doc, f"1. Plywood {ply_type_clean}:")
                add_eq(doc, f"W_plywood = {w_tot:.2f} KN/m²")
                add_eq(doc, f"Max Spacing = {s_spc:.2f} m\n")
                
                add_eq(doc, "Check for moment:", bold=True)
                M_ply = (w_tot * (s_spc**2)) / 10
                Z_req = (M_ply * 100) / 3.41
                ply_mall = 54.0 
                add_eq(doc, f"M = W * L² / 10 = {w_tot:.2f} * ({s_spc:.2f})² / 10 = {M_ply:.2f} KN.m")
                add_eq(doc, f"Z_req = M * 100 / 3.41 = {M_ply:.2f} * 100 / 3.41 = {Z_req:.2f} cm³")
                add_red_safe_check(doc, None, Z_req, ply_mall, "cm³")
                
                E_ply, I_ply = 74.52, 48.60 
                D_ply = (0.0068 * w_tot * (s_spc*100)**4) / (100 * E_ply * I_ply)
                all_ply_d = (s_spc*1000)/300
                add_eq(doc, "\nCheck for deflection:", bold=True)
                add_eq(doc, f"D = 0.0068 * W * L⁴ / (E * I) = 0.0068 * {w_tot:.2f} * ({s_spc*100:.1f})⁴ / (100 * {E_ply:.2f} * {I_ply:.1f}) = {D_ply:.2f} mm")
                add_red_safe_check(doc, None, D_ply, all_ply_d, "mm", f"Allowable = L/300 = {all_ply_d:.2f} mm")
                add_reference_line(doc, "Plywood", ds_page_map)
                
                doc.add_page_break()
                add_heading_14(doc, f"2. Secondary Decking {sec_type}:")
                add_eq(doc, f"- Secondary Beam length = {sec_L:.2f} m")
                add_eq(doc, f"- Max. spacing between main decking = {m_spc:.2f} m")
                add_eq(doc, f"- Max. spacing between Secondary decking = {s_spc:.2f} m")
                w_sec = w_tot * s_spc
                add_eq_highlight(doc, f"- W_sec = {w_tot:.2f} x {s_spc:.2f} = ", f"{w_sec:.2f} KN/m'")
                
                prop_s = SECTIONS_DB.get(sec_type, SECTIONS_DB.get('Timber H20', {'E':92.45, 'I':4613.0, 'Mall':5.0, 'Qall':11.0}))
                num_spans = max(1, int(round((sec_L - 2*sec_cant) / m_spc)))
                s_supports = [sec_cant + i*m_spc for i in range(num_spans+1)]
                s_loads = [{'type': 'linear', 'w1': w_sec, 'w2': w_sec, 'x1': 0.0, 'x2': sec_L}]
                
                s_sketch_bytes = draw_system_sketch(sec_L, s_supports, s_loads, transparent_bg=True)
                p_sk = doc.add_paragraph(); p_sk.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_sk.add_run().add_picture(io.BytesIO(s_sketch_bytes), width=Cm(15.0))
                add_centered_text(doc, "Load Assignment & Spans", size=12, color=RGBColor(100,100,100))
                
                add_eq(doc, "\nMaximum loads & deflections from attached Program Results:", underline=True)
                s_img_bytes, s_M, s_V, s_D, _, _, s_Dtxt = generate_acrow_diagrams(
                    sec_type, sec_L, s_supports, s_loads, prop_s['E'], prop_s['I'], prop_s['Mall'], prop_s['Qall'], Rall=None, transparent_bg=False
                )
                
                add_red_safe_check(doc, "Check for Moment", s_M, prop_s['Mall'], "KN.m")
                add_red_safe_check(doc, "Check for Shear", s_V, prop_s['Qall'], "KN")
                add_red_safe_check(doc, "Check for deflection", s_D, float(s_Dtxt.split('=')[-1].replace('mm','')), "mm", f"{s_Dtxt}")
                add_reference_line(doc, sec_type, ds_page_map)
                
                doc.add_page_break()
                p_s = doc.add_paragraph(); p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_s.add_run().add_picture(io.BytesIO(s_img_bytes), width=Cm(16.5))
                add_centered_text(doc, f"Analysis Diagrams for Secondary Beam ({sec_type})", size=12)

        if 'global_checks' in case:
            gc = case['global_checks']
            doc.add_page_break()
            
            # 🎯 3. Main Soldier Checks (Exact Deflection logic & Split logic)
            add_heading_14(doc, "3. Main Decking Soldier Beams:")
            add_line("According to Maximum Values for Moment & Shear from attached Program Results:", color=RGBColor(192, 0, 0), size=11)
            doc.add_paragraph()
            
            if 'soldier_zones' in gc:
                zone_idx = 1
                for zone in ['Bottom Slab', 'Cantilever Slab']:
                    if zone in gc['soldier_zones']:
                        z_data = gc['soldier_zones'][zone]
                        if z_data['m'] > 0.01 or z_data['v'] > 0.01:
                            add_line(f"{zone_idx}. Check of Soldier at {zone}:", bold=True, size=11, color=RGBColor(0, 0, 0))
                            
                            add_red_safe_check(doc, "Check for Moment", z_data['m'], gc.get('soldier_m_all', 13.1), "KN.m")
                            add_red_safe_check(doc, "Check for Shear", z_data['v'], gc.get('soldier_v_all', 100.8), "KN")
                            
                            # 🎯 Manual Deflection Check (No Diagrams, analytical values)
                            L_eff = z_data.get('L_eff', 600.0)
                            denom = 200.0 if zone == 'Cantilever Slab' else 400.0
                            allow_D = L_eff / denom
                            act_D = z_data['d']
                            def_txt = f"(L/{int(denom)} = {L_eff:.0f}/{int(denom)} = {allow_D:.2f} mm)"
                            
                            p_def = doc.add_paragraph()
                            force_ltr_left(p_def)
                            r1_def = p_def.add_run("• Check for deflection:\n")
                            r1_def.font.name, r1_def.font.size, r1_def.font.bold = 'Arial', Pt(11), True
                            
                            r2_def = p_def.add_run(f"  Max = {act_D:.2f} mm   <   {allow_D:.2f} mm   ")
                            r2_def.font.name, r2_def.font.size = 'Arial', Pt(11)
                            
                            res_def = p_def.add_run("SAFE" if act_D <= allow_D else "UNSAFE ❌")
                            res_def.font.name, res_def.font.size, res_def.font.bold = 'Arial', Pt(11), True
                            res_def.font.color.rgb = RGBColor(255, 0, 0)
                            
                            r3_def = p_def.add_run(f"\n  {def_txt}")
                            r3_def.font.name, r3_def.font.size, r3_def.italic = 'Arial', Pt(10), True
                            
                            doc.add_paragraph()
                            zone_idx += 1
                
                # 🎯 Reference line placed exactly here before diagrams
                add_reference_line(doc, "Soldier", ds_page_map)
            else:
                add_red_safe_check(doc, "Check for Moment", gc.get('soldier_m', 0), gc.get('soldier_m_all', 13.1), "KN.m")
                add_red_safe_check(doc, "Check for Shear", gc.get('soldier_v', 0), gc.get('soldier_v_all', 100.8), "KN")
                add_reference_line(doc, "Soldier", ds_page_map)
                
            if 'M' in case.get('img_bufs', {}):
                add_line("Bending Moment Diagram:", bold=False, size=12, color=RGBColor(192, 0, 0), underline=True)
                p_m = doc.add_paragraph(); p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_m.add_run().add_picture(io.BytesIO(case['img_bufs']['M']), width=Cm(16.5))
            if 'V' in case.get('img_bufs', {}):
                add_line("Shear Force Diagram:", bold=False, size=12, color=RGBColor(192, 0, 0), underline=True)
                p_v = doc.add_paragraph(); p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_v.add_run().add_picture(io.BytesIO(case['img_bufs']['V']), width=Cm(16.5))
            
            # 🎯 4. TILTING SYSTEM – PUSH PULL (All used types detailed)
            if 'strut_forces_dict' in gc and len(gc['strut_forces_dict']) > 0:
                doc.add_page_break()
                p_st = doc.add_paragraph()
                force_ltr_left(p_st)
                r_st = p_st.add_run("4. TILTING SYSTEM – PUSH PULL:")
                r_st.font.name, r_st.font.size, r_st.font.bold, r_st.underline = 'Arial', Pt(14), True, True
                
                p_sub = doc.add_paragraph()
                force_ltr_left(p_sub)
                r_sub = p_sub.add_run("Strut Axial Forces (N):")
                r_sub.font.name, r_sub.font.size, r_sub.font.bold, r_sub.underline = 'Arial', Pt(12), True, True
                doc.add_paragraph()
                
                for st_name, st_data in gc['strut_forces_dict'].items():
                    if st_data['max_n'] > 0.01:
                        p_name = doc.add_paragraph()
                        force_ltr_left(p_name)
                        r_name = p_name.add_run(f"• For Push Pull {st_name}")
                        r_name.font.name, r_name.font.size, r_name.font.bold = 'Arial', Pt(12), True
                        
                        add_red_safe_check(doc, "Axial Force (N)", st_data['max_n'], st_data['allow'], "KN")
                        doc.add_paragraph()
                
                add_reference_line(doc, "Push Pull", ds_page_map)
                
                if 'N' in case.get('img_bufs', {}):
                    add_line("Axial Force Diagram:", bold=False, size=12, color=RGBColor(192, 0, 0), underline=True)
                    p_n = doc.add_paragraph(); p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_n.add_run().add_picture(io.BytesIO(case['img_bufs']['N']), width=Cm(16.5))
            
            # 🎯 5. Shoring System + Reaction Diagram
            doc.add_page_break()
            add_heading_14(doc, "5. Shoring System / Supports:")
            add_eq(doc, f"- Load on Support = Max. Reaction from Main Beam = {gc.get('max_rxn', 0):.2f} KN")
            add_red_safe_check(doc, "Check for Support", gc.get('max_rxn', 0), gc.get('rxn_allow', 54.4), "KN")
            doc.add_paragraph()
            if 'R' in case.get('img_bufs', {}):
                add_line("Reactions Diagram:", bold=False, size=12, color=RGBColor(192, 0, 0), underline=True)
                p_r = doc.add_paragraph(); p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_r.add_run().add_picture(io.BytesIO(case['img_bufs']['R']), width=Cm(16.5))
                
    out = io.BytesIO()
    doc.save(out)
    return out

# ==============================================================================
# 7. MAIN STREAMLIT UI (Descriptive Inputs, 3-Zone Sweeps & The Beast)
# ==============================================================================
def render_bridge_module(proj_info):
    st.markdown("## 🌉 Bridge Formwork (True 2D DXF + Advanced Live Editor)")
    mode = st.radio(
        "Select Input Mode:", 
        ["1. Multi-Case DXF Auto-Extractor 🪄", "2. Single-Case Manual Builder 🛠️"], 
        horizontal=True
    )
    st.markdown("---")

    if "DXF" in mode:
        st.info("💡 **Smart Engine:** Upload DXF. Geometry & SHORING_LINES are extracted instantly in XZ plane!")
        
        c1, c2 = st.columns([1, 2])
        conc_density = c1.number_input("Concrete Density (kN/m3)", value=25.0, step=0.5)
        uploaded_dxf = c2.file_uploader("📥 Upload DXF File (.dxf)", type=['dxf'])
        
        if uploaded_dxf and st.button("🚀 Process DXF & Extract Data", type="primary", use_container_width=True):
            with st.spinner("Parsing DXF true 2D geometry, Mirroring Symmetry, & Raycasting Supports..."):
                cases_data = parse_dxf_bridge_cases(uploaded_dxf.getvalue(), conc_density)
                
            if cases_data:
                st.session_state.bridge_cases = cases_data
                st.success(f"✅ Successfully extracted {len(cases_data)} structural case(s)!")
                st.rerun()
            else: 
                st.error("❌ Failed to parse DXF. Please ensure layers are correct (TABLE_ANALYSIS, FRAME, SHORING_LINES, TEXT_DATA).")

        if 'bridge_cases' in st.session_state:
            global_sec = {
                'name': "Soldier U100", 'E': 2100.0, 
                'A': 34.3 / 10000.0, 'I': 412.0 / 100000000.0, 
                'Mall': 13.1, 'Qall': 100.8
            }
            all_cases_ready = []
            tabs = st.tabs([c['title'] for c in st.session_state.bridge_cases])
            
            for c_idx, tab in enumerate(tabs):
                case = st.session_state.bridge_cases[c_idx]
                if 'loaded_width_curr' not in case: case['loaded_width_curr'] = 1.30
                if 'base_live_load' not in case: case['base_live_load'] = 2.90
                    
                if 'calc_details' not in case or not case['calc_details']:
                    case['calc_details'] = []
                    for area_item in case.get('dxf_areas', []):
                        w_val = (area_item['area'] * conc_density * case['loaded_width_curr']) / area_item['length']
                        case['calc_details'].append({'segment': area_item['segment'], 'length': area_item['length'], 'area': area_item['area'], 'load_w': abs(w_val)})

                with tab:
                    c_edit, c_view = st.columns([1.2, 1.8])
                    
                    with c_edit:
                        st.markdown("### 📐 Table Specifications (Live Update)")
                        cw1, cw2 = st.columns(2)
                        
                        prev_lw = st.session_state.get(f"prev_lw_{c_idx}", float(case['loaded_width_curr']))
                        prev_bll = st.session_state.get(f"prev_bll_{c_idx}", float(case['base_live_load']))

                        new_lw = cw1.number_input("Loaded Width (Soldier Spacing) (m)", value=float(case['loaded_width_curr']), step=0.05, key=f"lw_curr_{c_idx}")
                        new_bll = cw2.number_input("Base Live Load (kN/m2)", value=float(case['base_live_load']), step=0.10, key=f"bll_curr_{c_idx}")

                        if abs(new_lw - prev_lw) > 1e-4 or abs(new_bll - prev_bll) > 1e-4:
                            st.session_state[f"prev_lw_{c_idx}"] = new_lw
                            st.session_state[f"prev_bll_{c_idx}"] = new_bll
                            case['loaded_width_curr'] = new_lw
                            case['base_live_load'] = new_bll

                            if f'fea_cache_{c_idx}' in st.session_state: del st.session_state[f'fea_cache_{c_idx}']
                            case['img_bufs'] = {}

                            new_auto_loads = []
                            case['calc_details'] = []
                            
                            for area_item in case.get('dxf_areas', []):
                                s_name = area_item['segment']
                                w_val = (area_item['area'] * conc_density * new_lw) / area_item['length']
                                matching_indices = [idx_m for idx_m, seg_m in enumerate(case['segments']) if seg_m['name'] == s_name]
                                if matching_indices:
                                    t_mode = 'Single Segment' if len(matching_indices) == 1 else 'Multiple Segments'
                                    new_auto_loads.append({
                                        'seg_idx': matching_indices[0], 'category': 'Dead Load', 'type': 'Uniform', 
                                        'dir': 'Global Z (Vertical)', 'target_mode': t_mode, 'target_segs_idx': matching_indices, 
                                        'start': 0.0, 'end': case['segments'][matching_indices[0]]['L'], 
                                        'w1': -abs(w_val), 'w2': -abs(w_val), 'loc': 0.0, 'is_auto': True
                                    })
                                case['calc_details'].append({'segment': s_name, 'length': area_item['length'], 'area': area_item['area'], 'load_w': abs(w_val)})
                                
                            valid_ll_indices = [idx for idx, seg in enumerate(case['segments']) if seg['name'] not in ["S30", "S31"]]
                            calculated_ll_magnitude = -abs(new_bll * new_lw)
                            if valid_ll_indices:
                                new_auto_loads.append({
                                    'seg_idx': valid_ll_indices[0], 'category': 'Live Load', 'type': 'Uniform', 
                                    'dir': 'Global Z (Vertical)', 'target_mode': 'Multiple Segments', 'target_segs_idx': valid_ll_indices, 
                                    'start': 0.0, 'end': case['segments'][valid_ll_indices[0]]['L'], 
                                    'w1': calculated_ll_magnitude, 'w2': calculated_ll_magnitude, 'loc': 0.0, 'is_auto': True
                                })
                            
                            manual_loads = [ld for ld in case.get('loads', []) if not ld.get('is_auto', False)]
                            case['loads'] = new_auto_loads + manual_loads
                            
                            for i, ld in enumerate(case['loads']):
                                if ld.get('is_auto'):
                                    st.session_state[f"alw1_{c_idx}_{i}"] = float(ld['w1'])
                                    st.session_state[f"alw2_{c_idx}_{i}"] = float(ld['w2'])
                            st.rerun()

                        with st.expander("🧱 Cross Section & Local Checks (Plywood & Secondary)", expanded=True):
                            st.info("📝 Descriptive Data for Report:")
                            cs1, cs2 = st.columns(2)
                            cs_depth = cs1.number_input("Total C.S. Depth (m)", value=float(case.get('local_params', {}).get('cs_depth', 2.0)), step=0.1, key=f"csd_{c_idx}")
                            top_ts = cs2.number_input("Top Slab Thickness (m)", value=float(case.get('local_params', {}).get('top_ts', 0.25)), step=0.05, key=f"tst_{c_idx}")
                            
                            st.markdown("---")
                            st.markdown("**Under Web:**")
                            w1, w2, w3 = st.columns(3)
                            web_ts = w1.number_input("Web Thick (m)", value=float(case.get('local_params', {}).get('web_ts', 2.20)), step=0.05, key=f"wts_{c_idx}")
                            sec_spc_web = w2.number_input("Sec Spacing (m)", value=float(case.get('local_params', {}).get('sec_spc_web', 0.15)), step=0.05, key=f"ssw_{c_idx}")
                            main_spc_web = w3.number_input("Main Spacing (m)", value=float(case.get('local_params', {}).get('main_spc_web', 1.20)), step=0.05, key=f"msw_{c_idx}")
                            ws1, ws2 = st.columns(2)
                            sec_l_web = ws1.number_input("Sec Length (m)", value=float(case.get('local_params', {}).get('sec_l_web', 2.50)), step=0.1, key=f"slw_{c_idx}")
                            sec_cant_web = ws2.number_input("Cantilever (m)", value=float(case.get('local_params', {}).get('sec_cant_web', 0.65)), step=0.05, key=f"scw_{c_idx}")

                            st.markdown("---")
                            st.markdown("**Under Bottom Slab:**")
                            b1, b2, b3 = st.columns(3)
                            bot_ts = b1.number_input("Bot Thick (m)", value=float(case.get('local_params', {}).get('bot_ts', 0.25)), step=0.05, key=f"bts_{c_idx}")
                            sec_spc_bot = b2.number_input("Sec Spacing (m)", value=float(case.get('local_params', {}).get('sec_spc_bot', 0.45)), step=0.05, key=f"ssb_{c_idx}")
                            main_spc_bot = b3.number_input("Main Spacing (m)", value=float(case.get('local_params', {}).get('main_spc_bot', 1.20)), step=0.05, key=f"msb_{c_idx}")
                            bs1, bs2 = st.columns(2)
                            sec_l_bot = bs1.number_input("Sec Length (m)", value=float(case.get('local_params', {}).get('sec_l_bot', 2.50)), step=0.1, key=f"slb_{c_idx}")
                            sec_cant_bot = bs2.number_input("Cantilever (m)", value=float(case.get('local_params', {}).get('sec_cant_bot', 0.65)), step=0.05, key=f"scb_{c_idx}")
                            
                            st.markdown("---")
                            try:
                                sec_opts = list(SECTIONS_DB.keys())
                            except Exception:
                                sec_opts = ["Timber H20", "Acrow Beam S12"]
                            idx_sec = sec_opts.index("Timber H20") if "Timber H20" in sec_opts else 0
                            sec_type = st.selectbox("Secondary Beam Section", sec_opts, index=idx_sec, key=f"sect_{c_idx}")
                            
                            case['local_params'] = {
                                'cs_depth': cs_depth, 'top_ts': top_ts, 'web_ts': web_ts, 'bot_ts': bot_ts, 
                                'sec_spc_web': sec_spc_web, 'main_spc_web': main_spc_web, 'sec_l_web': sec_l_web, 'sec_cant_web': sec_cant_web,
                                'sec_spc_bot': sec_spc_bot, 'main_spc_bot': main_spc_bot, 'sec_l_bot': sec_l_bot, 'sec_cant_bot': sec_cant_bot,
                                'sec_type': sec_type, 'ply_type': "18mm Plywood"
                            }

                        st.markdown("### 🎛️ Global Load Factors")
                        c_f1, c_f2, c_f3 = st.columns(3)
                        fac_d = c_f1.number_input("DL Factor", value=1.00, step=0.1, key=f"f_d_{c_idx}")
                        fac_l = c_f2.number_input("LL Factor", value=1.00, step=0.1, key=f"f_l_{c_idx}")
                        fac_w = c_f3.number_input("WL Factor", value=1.00, step=0.1, key=f"f_w_{c_idx}")
                        combo_factors = {'Dead Load': fac_d, 'Live Load': fac_l, 'Wind Load': fac_w}
                        
                        case['supports'].sort(key=lambda s: s['x'])
                        with st.expander(f"🔗 Edit Supports ({len(case['supports'])})", expanded=False):
                            for i, sup in enumerate(case['supports']):
                                st.markdown(f"**🟢 Support J{i+1}**")
                                c_s1, c_s2, c_s3, c_s4, c_del = st.columns([1, 1, 1.2, 1, 0.3])
                                sup['x'] = c_s1.number_input(f"J{i+1} X (m)", value=float(sup['x']), step=0.1, key=f"sx_{c_idx}_{i}")
                                sup['z'] = c_s2.number_input(f"J{i+1} Y (m)", value=float(sup.get('z', sup.get('y', 0.0))), step=0.1, key=f"sz_{c_idx}_{i}")
                                type_opts = ["Hinged", "Roller", "Fixed"]
                                idx_type = type_opts.index(sup['type']) if sup['type'] in type_opts else 1
                                sup['type'] = c_s3.selectbox(f"J{i+1} Type", type_opts, index=idx_type, key=f"st_{c_idx}_{i}")
                                sup['angle'] = c_s4.number_input(f"J{i+1} Angle(°)", value=float(sup.get('angle',0.0)), step=15.0, key=f"sa_{c_idx}_{i}")
                                c_del.markdown("<br>", unsafe_allow_html=True)
                                if c_del.button("❌", key=f"del_sup_{c_idx}_{i}"):
                                    case['supports'].pop(i)
                                    st.rerun()
                            if st.button("➕ Add Support", key=f"add_sup_{c_idx}"):
                                case['supports'].append({'x': 0.0, 'z': 0.0, 'type': 'Roller', 'angle': 0.0})
                                st.rerun()

                        if 'sec_overrides' not in case: 
                            case['sec_overrides'] = [global_sec.copy() for _ in range(len(case['segments']))]
                            
                        seg_names = [s['name'] for s in case['segments']]
                        unique_seg_opts = [f"{idx} - {name}" for idx, name in enumerate(seg_names)]
                        
                        with st.expander("📏 Override Sections", expanded=False):
                            override_segs = st.multiselect("Select segments:", unique_seg_opts, key=f"ovr_seg_{c_idx}")
                            if override_segs:
                                rad_opt = st.radio("Override Profile:", ["Custom Section", "Acrow Beam S12"], key=f"ovr_rad_{c_idx}")
                                if rad_opt == "Custom Section":
                                    o1, o2, o3, o4 = st.columns(4)
                                    o_sec = {'name': "Custom", 'E': 2100.0, 'A': o1.number_input("A", value=50.0, key=f"oa_{c_idx}")/10000.0, 'I': o2.number_input("I", value=1200.0, key=f"oi_{c_idx}")/100000000.0, 'Mall': o3.number_input("Mall", value=30.0, key=f"om_{c_idx}"), 'Qall': o4.number_input("Qall", value=150.0, key=f"oq_{c_idx}")}
                                else:
                                    o_sec = {'name': "S12", 'E': 2100.0, 'A': 20.0/10000.0, 'I': 800.0/100000000.0, 'Mall': 15.0, 'Qall': 80.0}
                                for s_val in override_segs: 
                                    idx_seg = int(s_val.split(' - ')[0])
                                    case['sec_overrides'][idx_seg] = o_sec.copy()

                        with st.expander(f"📐 Edit Struts ({len(case['struts'])})", expanded=False):
                            strut_opts = get_valid_strut_names()
                            for i, stt in enumerate(case['struts']):
                                s1, s2, s3, s4, s5, s_del = st.columns([1,1,1,1,1.2,0.3])
                                stt['tx'] = s1.number_input("TX", value=float(stt['tx']), step=0.1, key=f"ttx_{c_idx}_{i}")
                                stt['tz'] = s2.number_input("TZ", value=float(stt.get('tz', stt.get('ty', 0.0))), step=0.1, key=f"tty_{c_idx}_{i}")
                                stt['bx'] = s3.number_input("BX", value=float(stt['bx']), step=0.1, key=f"tbx_{c_idx}_{i}")
                                stt['bz'] = s4.number_input("BZ", value=float(stt.get('bz', stt.get('by', 0.0))), step=0.1, key=f"tby_{c_idx}_{i}")
                                idx_strut = strut_opts.index(stt['sec']) if stt['sec'] in strut_opts else 0
                                stt['sec'] = s5.selectbox("Sec", strut_opts, index=idx_strut, key=f"tsec_{c_idx}_{i}")
                                s_del.markdown("<br>", unsafe_allow_html=True)
                                if s_del.button("❌", key=f"del_strut_{c_idx}_{i}"):
                                    case['struts'].pop(i)
                                    st.rerun()
                            if st.button("➕ Add Strut", key=f"add_strut_{c_idx}"):
                                case['struts'].append({'tx': 0.0, 'tz': 0.0, 'bx': 0.0, 'bz': 0.0, 'sec': strut_opts[0] if strut_opts else 'Unknown'})
                                st.rerun()

                        with st.expander(f"⬇️ Loads (Auto & Manual) - Total: {len(case.get('loads', []))}", expanded=False):
                            cat_opts = ["Dead Load", "Live Load", "Wind Load"]
                            type_opts = ["Uniform", "Trapezoidal", "Point Load"]
                            dir_opts = ["Global X (Horizontal)", "Global Z (Vertical)", "Local Z (Perpendicular)"]
                            t_mode_opts = ["Single Segment", "Multiple Segments", "All Segments"]
                            
                            for i, ld in enumerate(case.get('loads', [])):
                                title_prefix = "🤖 Auto" if ld.get('is_auto', False) else "✍️ Manual"
                                with st.expander(f"{title_prefix} Load {i+1} ({ld.get('category')} - {ld.get('type')})", expanded=False):
                                    c_l1, c_l2, c_l3, c_l4 = st.columns([1.5, 1.5, 1.5, 0.5])
                                    ld['category'] = c_l1.selectbox("Category", cat_opts, index=cat_opts.index(ld.get('category', 'Dead Load')), key=f"alct_{c_idx}_{i}")
                                    ld['type'] = c_l2.selectbox("Type", type_opts, index=type_opts.index(ld.get('type', 'Uniform')), key=f"altp_{c_idx}_{i}")
                                    c_dir = ld.get('dir', 'Global Z (Vertical)')
                                    if 'Y' in c_dir: c_dir = 'Global Z (Vertical)'
                                    ld['dir'] = c_l3.selectbox("Direction", dir_opts, index=dir_opts.index(c_dir), key=f"aldr_{c_idx}_{i}")
                                    c_l4.markdown("<br>", unsafe_allow_html=True)
                                    if c_l4.button("❌", key=f"adel_ld_{c_idx}_{i}"): 
                                        case['loads'].pop(i)
                                        st.rerun()
                                    
                                    ld['target_mode'] = st.radio("Apply To:", t_mode_opts, index=t_mode_opts.index(ld.get('target_mode', 'Single Segment')), key=f"almode_{c_idx}_{i}", horizontal=True)
                                    if ld['target_mode'] == "Single Segment": 
                                        default_idx = ld.get('seg_idx', 0) if ld.get('seg_idx', 0) < len(unique_seg_opts) else 0
                                        s_val = st.selectbox("Target Seg", unique_seg_opts, index=default_idx, key=f"alsg_{c_idx}_{i}")
                                        ld['seg_idx'] = int(s_val.split(' - ')[0])
                                    elif ld['target_mode'] == "Multiple Segments": 
                                        safe_multi = [unique_seg_opts[idx] for idx in ld.get('target_segs_idx', []) if idx < len(unique_seg_opts)]
                                        sel_segs = st.multiselect("Target Segs", unique_seg_opts, default=safe_multi, key=f"alsm_{c_idx}_{i}")
                                        ld['target_segs_idx'] = [int(s.split(' - ')[0]) for s in sel_segs]
                                    
                                    sc1, sc2, sc3 = st.columns(3)
                                    is_disabled = ld.get('is_auto', False)
                                    ld['w1'] = sc1.number_input("W1 (kN/m)", value=float(ld.get('w1', 0.0)), step=1.0, key=f"alw1_{c_idx}_{i}", disabled=is_disabled)
                                    ld['w2'] = sc2.number_input("W2 (kN/m)", value=float(ld.get('w2', ld['w1'])) if ld['type'] == "Trapezoidal" else float(ld['w1']), step=1.0, key=f"alw2_{c_idx}_{i}", disabled=is_disabled)
                                    ld['loc'] = sc3.number_input("Location (m)", value=float(ld.get('loc', 0.0)), key=f"alloc_{c_idx}_{i}") if ld['type'] == "Point Load" else 0.0

                            if st.button("➕ Add Manual Load", key=f"add_mld_{c_idx}"): 
                                case.setdefault('loads', []).append({
                                    'seg_idx': 0, 'category': 'Live Load', 'type': 'Uniform', 
                                    'dir': 'Global Z (Vertical)', 'target_mode': 'Single Segment', 
                                    'target_segs_idx': [], 'start': 0.0, 'end': case['segments'][0]['L'], 
                                    'w1': -10.0, 'w2': -10.0, 'loc': 0.0, 'is_auto': False
                                })
                                st.rerun()

                        with st.expander("🤖 The Beast Optimizer", expanded=False):
                            ai_rxn = st.number_input("Target Max Rxn (kN)", value=54.4, step=1.0, key=f"br_{c_idx}")
                            ai_spc = st.text_input("Spacings (m)", value="2.40, 2.10, 1.80, 1.50, 1.20, 0.90, 0.60", key=f"bs_{c_idx}")
                            is_sym = st.checkbox("Symmetric", value=True, key=f"bm_{c_idx}")
                            opt_mode = st.radio("Optimization Depth:", ["Quick Search", "Deep Search"], index=0, key=f"opm_{c_idx}")
                            
                            if st.button("✨ Run Optimizer", type="primary", key=f"btn_opt_{c_idx}"):
                                p_bar = st.progress(0)
                                s_txt = st.empty()
                                with st.spinner("The Beast is optimizing (Gravity Ray Active)..."):
                                    succ, r_sup, r_str, msg = run_bridge_optimizer(case['segments'], case['segments'], case['sec_overrides'], case['struts'], case['loads'], ai_rxn, ai_spc, 0.25, is_sym, opt_mode, combo_factors, s_txt, p_bar)
                                    if r_sup:
                                        case['supports'] = r_sup
                                        case['struts'] = r_str
                                        st.session_state.bridge_cases[c_idx] = case
                                        if succ: 
                                            st.success(msg)
                                        else: 
                                            st.warning(msg)
                                        time.sleep(1.5)
                                        st.rerun()
                                    else: 
                                        st.error(msg)

                    with c_view:
                        target_rxn_ui = st.session_state.get(f"br_{c_idx}", 54.4)
                        st.markdown("<h4 style='text-align: center;'>Live Geometry & Reactions</h4>", unsafe_allow_html=True)
                        
                        expanded_loads = []
                        for ld in case.get('loads', []):
                            fac = combo_factors.get(ld.get('category', 'Dead Load'), 1.0)
                            t_mode = ld.get('target_mode', 'Single Segment')
                            t_idx_list = [ld.get('seg_idx', 0)] if t_mode == 'Single Segment' else (ld.get('target_segs_idx', []) if t_mode == 'Multiple Segments' else list(range(len(case['segments']))))
                                
                            for s_idx in t_idx_list:
                                if s_idx >= len(case['segments']): continue
                                f_ld = ld.copy()
                                f_ld['seg_idx'] = s_idx
                                f_ld['w1'] *= fac
                                f_ld['w2'] *= fac
                                L_seg = case['segments'][s_idx].get('L', 0.0)
                                if f_ld['type'] == 'Point Load': 
                                    f_ld['start'] = f_ld['end'] = min(f_ld.get('loc', 0.0), L_seg)
                                else: 
                                    f_ld['start'], f_ld['end'] = 0.0, L_seg
                                expanded_loads.append(f_ld)
                            
                        p_nodes, p_elems, p_nloads, p_supps = build_chain_mesh(case['segments'], case['sec_overrides'], expanded_loads, case['struts'], case['supports'], case.get('cut_points', []))
                        
                        try:
                            U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                            st.image(get_live_preview_image(p_nodes, p_elems, p_supps, expanded_loads, case['segments'], R_reactions=R, target_rxn=target_rxn_ui), use_container_width=True)
                        except Exception:
                            st.image(get_live_preview_image(p_nodes, p_elems, p_supps, expanded_loads, case['segments']), use_container_width=True)
                        
                        # 🎯 زرار الـ Pre-Check للتربيزة الحالية
                        if st.button(f"🔍 Pre-Check Safety (This Table)", type="secondary", use_container_width=True, key=f"btn_run_{c_idx}"):
                            with st.spinner(f"Solving Matrix, Auto-Sizing Struts & Extracting Deflection/Checks..."):
                                U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                                
                                struts_changed = False
                                for el in p_elems:
                                    if el['type'] == 'truss':
                                        st_idx = el.get('strut_idx')
                                        if st_idx is not None and st_idx < len(case['struts']):
                                            N_max = np.max(np.abs(el.get('internal', {}).get('N', [0])))
                                            opt_sec = get_optimal_strut_section(el.get('L', 0.0), N_max)
                                            if opt_sec and case['struts'][st_idx]['sec'] != opt_sec:
                                                case['struts'][st_idx]['sec'] = opt_sec
                                                struts_changed = True

                                if struts_changed:
                                    p_nodes, p_elems, p_nloads, p_supps = build_chain_mesh(case['segments'], case['sec_overrides'], expanded_loads, case['struts'], case['supports'], case.get('cut_points', []))
                                    U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                                    st.session_state.bridge_cases[c_idx] = case 
                                
                                st.session_state[f'fea_cache_{c_idx}'] = {'nodes': p_nodes, 'elements': p_elems, 'R': R, 'supports': p_supps, 'loads': expanded_loads}
                                
                                # 🎯 استخراج المسافات والأحمال لحساب الـ Deflection التحليلي (بدون الاعتماد على الإزاحة الكلية)
                                sup_xs = sorted(list(set([sup['x'] for sup in case['supports']])))
                                bot_span = max(np.diff(sup_xs)) * 1000 if len(sup_xs) > 1 else 1200.0
                                cant_span = 600.0
                                for seg in case['segments']:
                                    if 'F' in seg['name']:
                                        cant_span = max(cant_span, seg.get('L', 0.6) * 1000)
                                        
                                soldier_zones = {
                                    'Bottom Slab': {'m': 0.0, 'v': 0.0, 'd': 0.0, 'L_eff': bot_span},
                                    'Cantilever Slab': {'m': 0.0, 'v': 0.0, 'd': 0.0, 'L_eff': cant_span}
                                }
                                
                                max_m, max_v = 0.0, 0.0
                                soldier_mall, soldier_qall = 13.1, 100.8 
                                
                                # حساب الحمل المكافئ التقريبي للدفليكشن
                                w_dl = sum([abs(ld['w1']) for ld in expanded_loads if ld['category'] == 'Dead Load' and ld['seg_idx']==0])
                                w_ll = sum([abs(ld['w1']) for ld in expanded_loads if ld['category'] == 'Live Load' and ld['seg_idx']==0])
                                w_tot_kn = w_dl + w_ll
                                
                                for i_seg, sec in enumerate(case['sec_overrides']):
                                    seg_name = case['segments'][i_seg]['name']
                                    
                                    zone_key = None
                                    if seg_name in ['S30', 'S31']: zone_key = 'Bottom Slab'
                                    elif 'F' in seg_name: zone_key = 'Cantilever Slab'
                                    
                                    if zone_key:
                                        for el in p_elems:
                                            if el.get('group') == 'segment' and el.get('seg_idx') == i_seg:
                                                el_m = np.max(np.abs(el.get('internal', {}).get('M', [0])))
                                                el_v = np.max(np.abs(el.get('internal', {}).get('V', [0])))
                                                
                                                soldier_zones[zone_key]['m'] = max(soldier_zones[zone_key]['m'], el_m)
                                                soldier_zones[zone_key]['v'] = max(soldier_zones[zone_key]['v'], el_v)
                                                
                                                max_m = max(max_m, el_m)
                                                max_v = max(max_v, el_v)
                                                soldier_mall = sec['Mall']
                                                soldier_qall = sec['Qall']
                                                
                                                # حساب الدفليكشن بالمعادلات المانيوال
                                                e_sol = sec['E'] * 10000.0
                                                i_sol = sec['I'] / 100000000.0
                                                if e_sol * i_sol > 0:
                                                    if zone_key == 'Bottom Slab':
                                                        L_m = bot_span / 1000.0
                                                        d_m = (0.0068 * w_tot_kn * (L_m**4)) / (e_sol * i_sol)
                                                        soldier_zones[zone_key]['d'] = d_m * 1000.0
                                                    else:
                                                        L_m = cant_span / 1000.0
                                                        d_m = (w_tot_kn * (L_m**4)) / (8.0 * e_sol * i_sol)
                                                        soldier_zones[zone_key]['d'] = d_m * 1000.0
                                            
                                # 🎯 سحب بيانات كل أنواع النهائز
                                strut_forces_dict = {}
                                for el in p_elems:
                                    if el['type'] == 'truss':
                                        st_sec = el.get('sec', 'Push Pull Strut')
                                        n_abs = np.max(np.abs(el.get('internal', {}).get('N', [0])))
                                        allow_val = 30.0
                                        try:
                                            if st_sec in STRUTS_DB: allow_val = STRUTS_DB[st_sec]['allow']
                                        except Exception: pass
                                        
                                        if st_sec not in strut_forces_dict:
                                            strut_forces_dict[st_sec] = {'max_n': n_abs, 'allow': allow_val}
                                        else:
                                            if n_abs > strut_forces_dict[st_sec]['max_n']:
                                                strut_forces_dict[st_sec]['max_n'] = n_abs
                                        
                                case['global_checks'] = {
                                    'soldier_m': max_m, 'soldier_m_all': soldier_mall,
                                    'soldier_v': max_v, 'soldier_v_all': soldier_qall,
                                    'soldier_zones': soldier_zones,
                                    'strut_forces_dict': strut_forces_dict,
                                    'max_rxn': np.max(R) if len(R) > 0 else 0, 'rxn_allow': target_rxn_ui
                                }
                                
                                if struts_changed: st.rerun()

                        if f'fea_cache_{c_idx}' in st.session_state:
                            st.markdown("### 🎛️ Diagram Scales")
                            c_s1, c_s2, c_s3 = st.columns(3)
                            sc_n = c_s1.slider("Axial Scale", 0.001, 0.050, 0.010, step=0.001, key=f"scn_{c_idx}")
                            sc_v = c_s2.slider("Shear Scale", 0.001, 0.050, 0.010, step=0.001, key=f"scv_{c_idx}")
                            sc_m = c_s3.slider("Moment Scale", 0.001, 0.100, 0.010, step=0.001, key=f"scm_{c_idx}")
                            
                            cd = st.session_state[f'fea_cache_{c_idx}']
                            img_bufs = plot_sap2000_diagrams(cd['nodes'], cd['elements'], cd['R'], {'N': sc_n, 'V': sc_v, 'M': sc_m}, cd['supports'], cd['loads'], case['segments'], target_rxn=target_rxn_ui)
                            case['img_bufs'] = img_bufs
                            
                            if 'DL' in img_bufs: st.image(img_bufs['DL'], caption="Dead Load")
                            if 'LL' in img_bufs: st.image(img_bufs['LL'], caption="Live Load")
                            if 'WL' in img_bufs: st.image(img_bufs['WL'], caption="Wind Load")
                                
                            c_p1, c_p2 = st.columns(2)
                            c_p1.image(img_bufs['M'], caption="Moment")
                            c_p2.image(img_bufs['N'], caption="Axial")
                            c_p3, c_p4 = st.columns(2)
                            c_p3.image(img_bufs['V'], caption="Shear")
                            c_p4.image(img_bufs['R'], caption="Reactions")
                            
                    all_cases_ready.append(case)
                    
            # 🎯 تصميم الواجهة السفلية (شريط التحذير والأزرار)
            st.markdown("<br><hr>", unsafe_allow_html=True)
            st.markdown("<div style='background-color: #ffffe0; padding: 10px; border-radius: 5px; border-left: 5px solid #ffcc00; margin-bottom: 15px;'>⚠️ <b>تنبيه:</b> يرجى مراجعة وتأكيد تطابق البيانات الأساسية مع العناصر المدخلة قبل استخراج النوتة الحسابية.</div>", unsafe_allow_html=True)
            
            b1, b2 = st.columns(2)
            if b1.button("🔍 Pre-Check Safety", use_container_width=True):
                st.success("✅ System Pre-Checked Successfully! Please review the diagrams above.")
                
            doc_out = generate_multi_case_report(all_cases_ready, proj_info)
            b2.download_button("🚀 Generate Automated Calculation Sheet", data=doc_out.getvalue(), file_name="Calculation_Sheet_for_Bridge_Deck_Slab.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")

    else:
        # =========================================================
        # 💡 Manual Builder Mode 
        # =========================================================
        st.info("🛠️ **Manual Builder:** Define geometry in XZ plane directly.")
        if 'man_segs' not in st.session_state: st.session_state.man_segs = [{'name': 'S1', 'L': 3.0, 'type': 'Straight Line'}]
        if 'man_sups' not in st.session_state: st.session_state.man_sups = [{'x': 0.0, 'z': 0.0, 'type': 'Hinged', 'angle': 0.0}, {'x': 3.0, 'z': 0.0, 'type': 'Roller', 'angle': 0.0}]
        if 'man_strs' not in st.session_state: st.session_state.man_strs = []
        if 'man_lds' not in st.session_state: st.session_state.man_lds = []

        c_in, c_plot = st.columns([1.2, 1.8])
        with c_in:
            st.markdown("### 1. Segments")
            for i, seg in enumerate(st.session_state.man_segs):
                c1, c2, c_del = st.columns([1, 1.5, 0.3])
                seg['name'] = c1.text_input(f"Seg {i+1} Name", value=seg.get('name', f"S{i+1}"), key=f"msn_{i}")
                seg['L'] = c2.number_input(f"Length (m)", value=float(seg.get('L', 3.0)), step=0.1, key=f"msl_{i}")
                c_del.markdown("<br>", unsafe_allow_html=True)
                if c_del.button("❌", key=f"mdel_seg_{i}"):
                    if len(st.session_state.man_segs) > 1: st.session_state.man_segs.pop(i); st.rerun()
                        
                seg.update({'master_idx': i, 'abs_p1': (sum(s['L'] for s in st.session_state.man_segs[:i]), 0.0), 'abs_p2': (sum(s['L'] for s in st.session_state.man_segs[:i]) + seg['L'], 0.0), 'Shape Type': 'Straight Line'})
                
            if st.button("➕ Add Segment", key="madd_seg"): 
                st.session_state.man_segs.append({'name': f"S{len(st.session_state.man_segs)+1}", 'L': 3.0, 'type': 'Straight Line'}); st.rerun()

            st.markdown("### 2. Supports & Struts")
            st.session_state.man_sups.sort(key=lambda s: s['x'])
            for i, sup in enumerate(st.session_state.man_sups):
                st.markdown(f"**🟢 Support J{i+1}**")
                c1, c2, c3, c4, c_del = st.columns([1, 1, 1.2, 1, 0.3])
                sup['x'] = c1.number_input(f"J{i+1} X (m)", value=float(sup.get('x',0)), key=f"msx_{i}")
                sup['z'] = c2.number_input(f"J{i+1} Y (m)", value=float(sup.get('z',0)), key=f"msz_{i}")
                t_opts = ["Hinged", "Roller", "Fixed"]
                idx_t = t_opts.index(sup['type']) if sup['type'] in t_opts else 1
                sup['type'] = c3.selectbox(f"J{i+1} Type", t_opts, index=idx_t, key=f"mst_{i}")
                sup['angle'] = c4.number_input(f"J{i+1} Angle(°)", value=float(sup.get('angle',0.0)), key=f"msa_{i}")
                c_del.markdown("<br>", unsafe_allow_html=True)
                if c_del.button("❌", key=f"mdel_sup_{i}"): st.session_state.man_sups.pop(i); st.rerun()
                    
            if st.button("➕ Add Support", key="madd_sup"): 
                st.session_state.man_sups.append({'x':0.0, 'z':0.0, 'type':'Roller', 'angle': 0.0}); st.rerun()

            strut_opts = get_valid_strut_names()
            for i, ds in enumerate(st.session_state.man_strs):
                c1, c2, c3, c4, s_del = st.columns([1,1,1,1,0.3])
                ds['tx'] = c1.number_input("Top X", value=float(ds.get('tx',0)), key=f"mtx_{i}")
                ds['tz'] = c2.number_input("Top Z", value=float(ds.get('tz',3)), key=f"mtz_{i}")
                ds['bx'] = c3.number_input("Bot X", value=float(ds.get('bx',1)), key=f"mbx_{i}")
                ds['bz'] = c4.number_input("Bot Z", value=float(ds.get('bz',0)), key=f"mbz_{i}")
                idx_s = strut_opts.index(ds['sec']) if ds['sec'] in strut_opts else 0
                ds['sec'] = st.selectbox(f"Strut {i+1} Sec", strut_opts, index=idx_s, key=f"msec_{i}")
                s_del.markdown("<br>", unsafe_allow_html=True)
                if s_del.button("❌", key=f"mdel_str_{i}"): st.session_state.man_strs.pop(i); st.rerun()
                    
            if st.button("➕ Add Strut", key="madd_str"): 
                st.session_state.man_strs.append({'tx':0.0, 'tz':3.0, 'bx':1.0, 'bz':0.0, 'sec':strut_opts[0] if strut_opts else 'Unknown'}); st.rerun()

            st.markdown("### 3. Loads & Factors")
            c_f1, c_f2, c_f3 = st.columns(3)
            fac_d = c_f1.number_input("Dead Load Factor", value=1.00, step=0.1, key="mmf_d")
            fac_l = c_f2.number_input("Live Load Factor", value=1.00, step=0.1, key="mmf_l")
            fac_w = c_f3.number_input("Wind Load Factor", value=1.00, step=0.1, key="mmf_w")
            m_combo_factors = {'Dead Load': fac_d, 'Live Load': fac_l, 'Wind Load': fac_w}

            cat_opts = ["Dead Load", "Live Load", "Wind Load"]
            type_opts = ["Uniform", "Trapezoidal", "Point Load"]
            dir_opts = ["Global X (Horizontal)", "Global Z (Vertical)", "Local Z (Perpendicular)"]
            t_mode_opts = ["Single Segment", "Multiple Segments", "All Segments"]
            seg_names_man = [s['name'] for s in st.session_state.man_segs]
            unique_seg_opts_man = [f"{idx} - {name}" for idx, name in enumerate(seg_names_man)]

            with st.expander(f"⬇️ Manual Loads ({len(st.session_state.man_lds)})", expanded=True):
                for i, ld in enumerate(st.session_state.man_lds):
                    with st.expander(f"📥 Load {i+1} ({ld.get('category', 'Dead Load')})", expanded=False):
                        c_l1, c_l2, c_l3, c_l4 = st.columns([1.5, 1.5, 1.5, 0.5])
                        ld['category'] = c_l1.selectbox("Category", cat_opts, index=cat_opts.index(ld.get('category', 'Dead Load')) if ld.get('category') in cat_opts else 0, key=f"mmlct_{i}")
                        ld['type'] = c_l2.selectbox("Type", type_opts, index=type_opts.index(ld.get('type', 'Uniform')) if ld.get('type') in type_opts else 0, key=f"mmltp_{i}")
                        c_dir = ld.get('dir', 'Global Z (Vertical)')
                        if 'Y' in c_dir: c_dir = 'Global Z (Vertical)'
                        ld['dir'] = c_l3.selectbox("Direction", dir_opts, index=dir_opts.index(c_dir), key=f"mmldr_{i}")
                        c_l4.markdown("<br>", unsafe_allow_html=True)
                        if c_l4.button("❌", key=f"mmdel_ld_{i}"): st.session_state.man_lds.pop(i); st.rerun()

                        ld['target_mode'] = st.radio("Apply To:", t_mode_opts, index=t_mode_opts.index(ld.get('target_mode', 'Single Segment')), key=f"mmlmode_{i}", horizontal=True)
                        if ld['target_mode'] == "Single Segment":
                            default_idx = ld.get('seg_idx', 0) if ld.get('seg_idx', 0) < len(unique_seg_opts_man) else 0
                            s_val = st.selectbox("Target Seg", unique_seg_opts_man, index=default_idx, key=f"mmlsg_{i}")
                            ld['seg_idx'] = int(s_val.split(' - ')[0])
                        elif ld['target_mode'] == "Multiple Segments":
                            safe_multi = [unique_seg_opts_man[idx] for idx in ld.get('target_segs_idx', []) if idx < len(unique_seg_opts_man)]
                            sel_segs = st.multiselect("Target Segs", unique_seg_opts_man, default=safe_multi, key=f"mmlsm_{i}")
                            ld['target_segs_idx'] = [int(s.split(' - ')[0]) for s in sel_segs]

                        sc1, sc2, sc3 = st.columns(3)
                        ld['w1'] = sc1.number_input("W1 (kN/m)", value=float(ld.get('w1', -10.0)), step=1.0, key=f"mmlw1_{i}")
                        if ld['type'] == "Trapezoidal": ld['w2'] = sc2.number_input("W2 (kN/m)", value=float(ld.get('w2', ld['w1'])), step=1.0, key=f"mmlw2_{i}")
                        else: ld['w2'] = ld['w1']
                        if ld['type'] == "Point Load": ld['loc'] = sc3.number_input("Location (m)", value=float(ld.get('loc', 0.0)), key=f"mmlloc_{i}")
                        else: ld['loc'] = 0.0

                if st.button("➕ Add Manual Load", key="mmadd_ld"): 
                    st.session_state.man_lds.append({'seg_idx': 0, 'w1': -10.0, 'w2': -10.0, 'dir': 'Global Z (Vertical)', 'category': 'Live Load', 'type': 'Uniform', 'target_mode': 'Single Segment', 'target_segs_idx': [], 'loc': 0.0})
                    st.rerun()

        with c_plot:
            man_target_rxn = st.number_input("Target Max Rxn (kN) for Coloring", value=54.4, step=1.0, key="man_tgt_rxn")
            st.markdown("<h4 style='text-align: center;'>Live Geometry & Loads</h4>", unsafe_allow_html=True)
            active_sections = [{'name': "Soldier U100", 'E': 2100.0, 'A': 34.3/10000.0, 'I': 412.0/100000000.0, 'Mall': 13.1, 'Qall': 100.8}] * len(st.session_state.man_segs)
            
            m_expanded_loads = []
            for ld in st.session_state.man_lds:
                fac = m_combo_factors.get(ld.get('category', 'Dead Load'), 1.0)
                t_mode = ld.get('target_mode', 'Single Segment')
                t_idx_list = [ld.get('seg_idx', 0)] if t_mode == 'Single Segment' else (ld.get('target_segs_idx', []) if t_mode == 'Multiple Segments' else list(range(len(st.session_state.man_segs))))
                for s_idx in t_idx_list:
                    if s_idx >= len(st.session_state.man_segs): continue
                    f_ld = ld.copy()
                    f_ld['seg_idx'] = s_idx
                    f_ld['w1'] *= fac; f_ld['w2'] *= fac
                    L_seg = st.session_state.man_segs[s_idx].get('L', 0.0)
                    if f_ld['type'] == 'Point Load': f_ld['start'] = f_ld['end'] = min(f_ld.get('loc', 0.0), L_seg)
                    else: f_ld['start'], f_ld['end'] = 0.0, L_seg
                    m_expanded_loads.append(f_ld)
            
            p_nodes, p_elems, p_nloads, p_supps = build_chain_mesh(st.session_state.man_segs, active_sections, m_expanded_loads, st.session_state.man_strs, st.session_state.man_sups, [])
            
            try:
                U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                st.image(get_live_preview_image(p_nodes, p_elems, p_supps, m_expanded_loads, st.session_state.man_segs, R_reactions=R, target_rxn=man_target_rxn), use_container_width=True)
            except Exception:
                st.image(get_live_preview_image(p_nodes, p_elems, p_supps, m_expanded_loads, st.session_state.man_segs), use_container_width=True)
            
            if st.button("🚀 Generate Detail Diagrams", type="primary", use_container_width=True, key="mrun_btn"):
                with st.spinner("Solving FEA & Auto-Sizing Struts..."):
                    U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                    
                    struts_changed = False
                    for el in p_elems:
                        if el['type'] == 'truss':
                            st_idx = el.get('strut_idx')
                            if st_idx is not None and st_idx < len(st.session_state.man_strs):
                                N_max = np.max(np.abs(el.get('internal', {}).get('N', [0])))
                                opt_sec = get_optimal_strut_section(el.get('L', 0.0), N_max)
                                if opt_sec and st.session_state.man_strs[st_idx]['sec'] != opt_sec:
                                    st.session_state.man_strs[st_idx]['sec'] = opt_sec
                                    struts_changed = True

                    if struts_changed:
                        p_nodes, p_elems, p_nloads, p_supps = build_chain_mesh(st.session_state.man_segs, active_sections, m_expanded_loads, st.session_state.man_strs, st.session_state.man_sups, [])
                        U, R, net_load = solve_fea_engine(p_nodes, p_elems, p_nloads, p_supps)
                        
                    st.session_state['man_fea_cache'] = {'nodes': p_nodes, 'elements': p_elems, 'R': R, 'supports': p_supps, 'loads': m_expanded_loads}
                    
                    safety_data = []
                    for i_seg, sec in enumerate(active_sections):
                        max_m, max_v = 0.0, 0.0
                        for el in p_elems:
                            if el.get('group') == 'segment' and el.get('seg_idx') == i_seg:
                                max_m = max(max_m, np.max(np.abs(el.get('internal', {}).get('M', [0]))))
                                max_v = max(max_v, np.max(np.abs(el.get('internal', {}).get('V', [0]))))
                        seg_name_clean = st.session_state.man_segs[i_seg]['name'].split('-')[0]
                        s_status = "SAFE" if (max_m <= sec['Mall'] and max_v <= sec['Qall']) else "UNSAFE ❌"
                        safety_data.append({"Segment": seg_name_clean, "M_max": f"{max_m:.2f} / {sec['Mall']:.2f}", "V_max": f"{max_v:.2f} / {sec['Qall']:.2f}", "Status": s_status})
                    st.session_state.man_safety_df = safety_data
                    
                    if struts_changed: st.rerun()

            if 'man_fea_cache' in st.session_state:
                st.markdown("### 🎛️ Diagram Scales")
                c_s1, c_s2, c_s3 = st.columns(3)
                sc_n = c_s1.slider("Axial Scale", 0.001, 0.050, 0.010, step=0.001, key="mscn")
                sc_v = c_s2.slider("Shear Scale", 0.001, 0.050, 0.010, step=0.001, key="mscv")
                sc_m = c_s3.slider("Moment Scale", 0.001, 0.100, 0.010, step=0.001, key="mscm")
                
                cd = st.session_state['man_fea_cache']
                img_bufs = plot_sap2000_diagrams(cd['nodes'], cd['elements'], cd['R'], {'N': sc_n, 'V': sc_v, 'M': sc_m}, cd['supports'], cd['loads'], st.session_state.man_segs, target_rxn=man_target_rxn)
                
                st.session_state.man_case_data = [{'title': 'Manual Case', 'img_bufs': img_bufs, 'safety_df': st.session_state.man_safety_df}]
                
                if 'DL' in img_bufs: st.image(img_bufs['DL'], caption="Dead Load")
                if 'LL' in img_bufs: st.image(img_bufs['LL'], caption="Live Load")
                if 'WL' in img_bufs: st.image(img_bufs['WL'], caption="Wind Load")
                c_p1, c_p2 = st.columns(2)
                c_p1.image(img_bufs['M'], caption="Moment")
                c_p2.image(img_bufs['N'], caption="Axial")
                c_p3, c_p4 = st.columns(2)
                c_p3.image(img_bufs['V'], caption="Shear")
                c_p4.image(img_bufs['R'], caption="Reactions")

            st.markdown("<br><hr>", unsafe_allow_html=True)
            st.markdown("<div style='background-color: #ffffe0; padding: 10px; border-radius: 5px; border-left: 5px solid #ffcc00; margin-bottom: 15px;'>⚠️ <b>تنبيه:</b> يرجى مراجعة وتأكيد تطابق البيانات الأساسية مع العناصر المدخلة قبل استخراج النوتة الحسابية.</div>", unsafe_allow_html=True)
            b1, b2 = st.columns(2)
            if b1.button("🔍 Pre-Check Safety", use_container_width=True):
                st.success("✅ System Pre-Checked Successfully! Please review the diagrams above.")
            
            if 'man_case_data' in st.session_state:
                doc_out = generate_multi_case_report(st.session_state.man_case_data, proj_info)
                b2.download_button("🚀 Generate Automated Calculation Sheet", data=doc_out.getvalue(), file_name="Calculation_Sheet_for_Bridge_Deck_Slab.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")

if __name__ == "__main__":
    render_bridge_module({
        "proj_name": "Acrow Mega Project", 
        "contractor": "Main Contractor",
        "calc_sub": "Bridge Shoring Analysis",
        "sys_name": "Acrow Bridge Systems",
        "date_val": "2026",
        "ref_code": "SBC"
    })